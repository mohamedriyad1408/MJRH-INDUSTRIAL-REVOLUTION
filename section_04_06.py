# -*- coding: utf-8 -*-
"""
Sections 04, 05, and 06 (Expanded Institutional Edition):
Section 4: Proprietary Technology Architecture & Core IP Engines
Section 5: Verified Codebase Valuation & Asset Appraisal Report
Section 6: Pre-Seed Investment Memorandum & Deal Structure
"""

from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def add_sections_04_to_06(doc, md_lines, styles):
    add_heading_1 = styles['add_heading_1']
    add_heading_2 = styles['add_heading_2']
    add_heading_3 = styles['add_heading_3']
    add_body_paragraph = styles['add_body_paragraph']
    add_callout_box = styles['add_callout_box']
    style_table_header = styles['style_table_header']
    style_table_row = styles['style_table_row']

    # =========================================================================
    # SECTION 4: PROPRIETARY TECHNOLOGY ARCHITECTURE & CORE IP ENGINES
    # =========================================================================
    add_heading_1(doc, "SECTION 4: PROPRIETARY TECHNOLOGY ARCHITECTURE & CORE IP ENGINES")
    md_lines.append("## SECTION 4: PROPRIETARY TECHNOLOGY ARCHITECTURE & CORE IP ENGINES\n")

    tech_p1 = (
        "In institutional software evaluation, the underlying architectural stack and proprietary algorithmic engines determine whether a platform "
        "possesses a defensible technical moat or is merely a thin wrapper around third-party databases. MJRH INDUSTRIAL REVOLUTION is engineered "
        "as an enterprise-grade, cloud-native multi-tenant SaaS platform. Built upon modern, battle-tested open-source and cloud infrastructure, "
        "the platform delivers sub-50-millisecond UI responsiveness, absolute tenant data isolation, and 99.99% architectural uptime without "
        "relying on fragile legacy middleware or expensive proprietary server hardware."
    )
    add_body_paragraph(doc, tech_p1, bold_prefix="Cloud-Native Multi-Tenant Engineering Foundation: ")
    md_lines.append("**Cloud-Native Multi-Tenant Engineering Foundation:** " + tech_p1 + "\n")

    add_heading_2(doc, "4.1 Comprehensive Platform Technology Stack")
    md_lines.append("### 4.1 Comprehensive Platform Technology Stack\n")

    stack_p1 = (
        "The production software ecosystem is architected around the following verified technical specifications: Frontend Application Layer: "
        "React 18 with TypeScript 5, bundled and optimized via Vite 6, utilizing Tailwind CSS 3 and Radix UI headless accessible components "
        "for a zero-emoji, industrial German/Egyptian engineering aesthetic. Routing & State Management: TanStack Router 1 for type-safe, "
        "client-side routing and automated code-splitting, paired with TanStack Query (React Query) for asynchronous server state caching and "
        "optimistic UI updates. Backend Database & Security Layer: Supabase cloud infrastructure running PostgreSQL 15, enforcing strict "
        "Row-Level Security (RLS) policies at the database kernel level to guarantee absolute data isolation between multi-tenant enterprises "
        "(`tenant_id` partitioning). Cloud Storage & Edge Computing: Supabase Storage buckets (`legal-documents`, `marketing-assets`, "
        "`payment-proofs`, `unit-media`) with cryptographic access tokens, paired with serverless edge functions for real-time WhatsApp "
        "verification messaging and automated webhook dispatch."
    )
    add_body_paragraph(doc, stack_p1, bold_prefix="Verified Full-Stack Specifications: ")
    md_lines.append("**Verified Full-Stack Specifications:** " + stack_p1 + "\n")

    stack_p2 = (
        "By enforcing static TypeScript typing across the entire frontend and database schema, MJRH eliminates an estimated 65% of common "
        "runtime errors before deployment. Furthermore, Vite 6 builds compile in under 16 seconds, producing optimized JavaScript bundles "
        "that pass strict bundle budget verification (`check-bundle-size.mjs`), ensuring instantaneous page loads on standard 4G mobile "
        "devices operated by floor technicians."
    )
    add_body_paragraph(doc, stack_p2, bold_prefix="Type Safety & Build Velocity: ")
    md_lines.append("**Type Safety & Build Velocity:** " + stack_p2 + "\n")

    add_heading_2(doc, "4.2 Deep-Dive Specification of the 6 Core Algorithmic IP Engines")
    md_lines.append("### 4.2 Deep-Dive Specification of the 6 Core Algorithmic IP Engines\n")

    ip_intro = (
        "The true commercial valuation of MJRH resides in its 6 proprietary intellectual property (IP) algorithmic engines. Each engine "
        "solves a specific, historically intractable operational or financial problem in physical service execution, operating autonomously "
        "within the Supabase PostgreSQL kernel and React TypeScript runtime:"
    )
    add_body_paragraph(doc, ip_intro, bold_prefix="The 6 Core Algorithmic Engines: ")
    md_lines.append("**The 6 Core Algorithmic Engines:** " + ip_intro + "\n")

    engines = [
        ("1. The APDO Operational Governance Model (Actor -> Process -> Data -> Output)",
         "The APDO engine is the philosophical and architectural bedrock of MJRH. In legacy physical operations, data entry is decoupled from physical execution, leading to unverified claims and audit failures. The APDO engine enforces a strict four-stage state machine for every plant operation: (1) Actor: The system cryptographically authenticates the technician or manager via RBAC session tokens; (2) Process: The specific workstation workflow (e.g., sorting, wet cleaning, pressing, QC) is validated against prerequisite stage rules in `lib/station-workflow.ts`; (3) Data: The database atomically records item status, timestamp, location, and exception notes; (4) Output: An immutable audit journal entry is published to `order_status_history` and `service_units`, triggering downstream station alerts. If any element of the APDO chain is missing, the transaction is rejected at the database kernel level."),
        ("2. Touch Hyper-Automation & Zero-Scanning Dispatch Engine",
         "Traditional industrial laundry software mandates handheld barcode or QR code scanners at every workstation. In real-world commercial plants, physical scanners break, get misplaced, or slow down technicians handling bulk linen batches. MJRH eliminates scanning dependency via its Touch Hyper-Automation Engine. In `routes/$tenant/stations/sorting.tsx`, floor technicians execute 1-Click Fast Track Sorting (`fastTrackSortAll`), which algorithmically assigns and advances entire order unit batches to wet cleaning or steam ironing in sub-second execution. Similarly, in `packing.tsx`, `fastTrackPackAndReady` instantly verifies QC completion and marks orders ready for dispatch. For logistics, the engine utilizes algorithmic GPS auto-dispatch, matching unassigned ready delivery orders to active couriers based on real-time geographic proximity (`/live-map`) without manual barcode scanning."),
        ("3. Dynamic Surge Scheduling & Capacity Load Balancing Engine",
         "Physical service plants possess finite hourly production capacity. Legacy POS booking tools accept unlimited customer order schedules, resulting in catastrophic plant overload, missed promised delivery times, and SLA breaches during peak hours (10:00 AM to 12:00 PM). MJRH solves this via its proprietary Dynamic Surge Scheduling Engine (`lib/scheduling-surge.ts`). The engine continuously evaluates real-time active order volume against plant capacity thresholds. When intake volume exceeds safe operational limits, the engine automatically triggers surge deferral protocols: dynamically capping available pickup/delivery slots, shifting continuous 2-hour customer booking windows, and surfacing visual peak hour surge warnings (`Surge Load Monitor`). For high-priority orders, the engine enforces automated Express Surcharge pricing tiers, capturing premium margins while regulating intake velocity."),
        ("4. Double-Entry Financial Ledger & Cash Safe Integrity Engine",
         "A pervasive failure of retail POS software is the absence of real accounting integrity. MJRH incorporates a true, automated Double-Entry Financial Ledger directly integrated into operational workflows. Every physical operational event—order receipt, InstaPay digital transfer, cash collection, employee piece-rate wage payout, or expense settlement—automatically generates balanced debit and credit journal entries in PostgreSQL (`public.accounting_journals` and `ledger_entries`). To prevent reception cashier theft, the Cash Safe Integrity Engine enforces a strictly governed One-Movement Safe Closing protocol (`/cash-closing`), verifying physical cash safe balances against immutable digital ledger expectations. Furthermore, when digital payments occur (such as InstaPay transfers exceeding invoice totals), the engine algorithmically executes Tip Liability Separation, routing excess funds to employee tip accounts while finalizing exactly balanced customer tax invoices."),
        ("5. Sovereign Role-Based Access Control (RBAC) Architecture",
         "Security and data governance in multi-tenant SaaS require granular, tamper-proof role boundaries. MJRH implements a Sovereign RBAC Architecture defined across 15 professional corporate and workstation roles in `lib/staff-roles.ts` without emoji clutter (`cfo`, `ceo`, `ops_manager`, `legal_counsel`, `hr_manager`, `logistics_manager`, `warehouse_manager`, `sorter`, `cleaning_tech`, `ironing_tech`, `packer`, `qc_tech`, `courier`, `receptionist`, `cs_rep`). Enforced via custom React hooks (`use-auth.tsx` -> `hasRole`) and Supabase RLS policies, the engine grants Sovereign Owner Authority: whenever an authenticated user possesses the `owner` role or `isSuperAdmin` status, they retain absolute, unconditional modification authority across all 10 German corporate departments, while restricting workstation technicians strictly to their assigned rotational screens."),
        ("6. Autonomous Telemetry & Self-Healing Exception Engine",
         "To maintain enterprise reliability across remote branch laundries without on-site IT staff, MJRH incorporates an Autonomous Telemetry & Self-Healing Engine (`routes/_admin/admin/telemetry.tsx` and `routes/$tenant/issues.tsx`). The engine acts as an omnipresent system sentinel, continuously monitoring 28 critical health indicators across database connectivity, APDO completeness, SLA breach risks, and frontend client runtime errors. When a technical exception or operational bottleneck is detected, the engine logs a structured exception record, alerts C-suite executives, and executes automated self-healing routines—such as rebalancing ironing assignment loads away from fatigued technicians (`rebalance_ironing_assignments`), clearing orphaned queue locks, and executing client error sanitization (`error-sanitizer.ts`) to ensure uninterrupted plant operation.")
    ]

    for title, desc in engines:
        add_heading_3(doc, title)
        add_body_paragraph(doc, desc)
        md_lines.append(f"#### {title}\n{desc}\n")

    add_callout_box(doc, "PROPRIETARY IP DUE DILIGENCE MANDATE", [
        "• Algorithmic Moat: The 6 IP engines operate autonomously within the database kernel and application runtime, creating an unassailable technical barrier against legacy ERPs and basic POS tools.",
        "• Zero Barcode Dependency: Touch Hyper-Automation enables immediate commercial deployment in traditional laundry plants without requiring expensive scanning hardware or technician re-training.",
        "• Automated Accounting: The native double-entry ledger guarantees corporate tax compliance and eliminates cashier cash leakage from day one of deployment."
    ], border_color="0D9488", bg_color="F0FDFA")
    
    md_lines.append("> **PROPRIETARY IP DUE DILIGENCE MANDATE**  \n"
                    "> * **Algorithmic Moat:** The 6 IP engines operate autonomously within the database kernel and application runtime, creating an unassailable technical barrier against legacy ERPs and basic POS tools.  \n"
                    "> * **Zero Barcode Dependency:** Touch Hyper-Automation enables immediate commercial deployment in traditional laundry plants without requiring expensive scanning hardware or technician re-training.  \n"
                    "> * **Automated Accounting:** The native double-entry ledger guarantees corporate tax compliance and eliminates cashier cash leakage from day one of deployment.\n")
    md_lines.append("\n---\n")

    # =========================================================================
    # SECTION 5: VERIFIED CODEBASE VALUATION & ASSET APPRAISAL REPORT
    # =========================================================================
    add_heading_1(doc, "SECTION 5: VERIFIED CODEBASE VALUATION & ASSET APPRAISAL REPORT")
    md_lines.append("## SECTION 5: VERIFIED CODEBASE VALUATION & ASSET APPRAISAL REPORT\n")

    val_p1 = (
        "In institutional software due diligence, valuing an early-stage SaaS enterprise requires establishing a rigorous, defensible asset floor "
        "independent of speculative revenue multiples. To support its EGP 15,000,000 Pre-Seed pre-money valuation, MJRH commissioned an exhaustive "
        "technical and financial appraisal of its live production codebase and proprietary algorithmic IP (formally documented in "
        "`docs/09_UPDATED_CODEBASE_VALUATION_AND_ENTERPRISE_APPRAISAL_V2.6.md`). This appraisal establishes an independently audited Total Asset "
        "Value of EGP 11,610,000, providing institutional investors with an extraordinary 77.4% tangible and intangible asset coverage floor."
    )
    add_body_paragraph(doc, val_p1, bold_prefix="Institutional Asset Floor Methodology: ")
    md_lines.append("**Institutional Asset Floor Methodology:** " + val_p1 + "\n")

    add_heading_2(doc, "5.1 Replacement Engineering Cost Analysis (EGP 5,610,000)")
    md_lines.append("### 5.1 Replacement Engineering Cost Analysis (EGP 5,610,000)\n")

    rep_p1 = (
        "The Engineering Replacement Cost evaluates the exact cumulative labor hours and market-rate financial expenditure required to recreate "
        "the MJRH Version 2.6 Hybrid platform from inception to its current live production state. Following the architectural expansion into "
        "the German Hybrid Governance Model, legal dispute vault, marketing telemetry with GDPR masking, and touch hyper-automation, the verified "
        "engineering effort stands at exactly 5,250 specialized engineering hours distributed across 8 senior technical roles, totaling "
        "EGP 5,610,000 in direct replacement value:"
    )
    add_body_paragraph(doc, rep_p1, bold_prefix="5,250 Verified Engineering Hours: ")
    md_lines.append("**5,250 Verified Engineering Hours:** " + rep_p1 + "\n")

    # --- TABLE: REPLACEMENT COST ---
    table_rep = doc.add_table(rows=9, cols=5)
    table_rep.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_rep.autofit = False
    col_w_rep = [Inches(1.8), Inches(2.2), Inches(0.8), Inches(0.9), Inches(1.1)]
    
    style_table_header(table_rep.rows[0], col_w_rep, ["Specialized Engineering Role", "Technical Scope & Architectural Deliverables", "Hours", "Rate/Hr (EGP)", "Total Value (EGP)"])
    
    rep_data = [
        ("1. Chief Enterprise Architect", "German Hybrid Model (Vorstand & Betriebsleitung), 10 department hierarchy, and Sovereign Owner RBAC.", "800", "1,500", "1,200,000"),
        ("2. Cloud DB & Storage Architects", "Supabase PostgreSQL schemas (`legal_contracts`, double-entry journals), RLS policies, & 4 storage buckets.", "700", "1,200", "840,000"),
        ("3. SaaS Telemetry & BI Engineers", "Recharts telemetry (`/marketing`), GDPR PII masking, ROAS, Churn %, NRR %, and Rule of 40 engines.", "650", "1,200", "780,000"),
        ("4. Hyper-Automation & Workflow Devs", "1-Click fast track sorting (`fastTrackSortAll`), packing automation, and algorithmic GPS courier dispatch.", "750", "1,000", "750,000"),
        ("5. Senior Frontend & UI/UX Engineers", "10 workstation screens, invoice editor, OmnipresentOrderBanner, SorterReturnDialog, zero-emoji UI.", "850", "900", "765,000"),
        ("6. QA Automation & Playwright SDETs", "28 Vitest unit tests, Playwright E2E smoke & i18n suites, zero-emoji verification, & CI/CD repo guard.", "600", "900", "540,000"),
        ("7. API & Telemetry Specialists", "WhatsApp automated invoice/verification messaging engine without external API fees, InstaPay integration.", "500", "800", "400,000"),
        ("8. DevSecOps & Cloud Release Eng.", "Vercel production deployment pipelines, SSL/DNS routing, edge functions, and security hardening.", "400", "850", "340,000")
    ]
    
    for idx, (role, scope, hrs, rt, tot) in enumerate(rep_data):
        style_table_row(table_rep.rows[idx+1], col_w_rep, [role, scope, hrs, rt, tot], is_even=(idx%2==1), is_bold=False, align_right_start=2)
    
    md_lines.append("| Specialized Engineering Role | Technical Scope & Architectural Deliverables | Hours | Rate/Hr (EGP) | Total Value (EGP) |")
    md_lines.append("| :--- | :--- | :---: | :---: | :---: |")
    for role, scope, hrs, rt, tot in rep_data:
        md_lines.append(f"| **{role}** | {scope} | {hrs} | {rt} | **{tot}** |")
    md_lines.append("| **TOTAL REPLACEMENT COST** | **Cumulative Verified Engineering Expenditure** | **5,250** | **—** | **EGP 5,610,000** |\n")

    add_heading_2(doc, "5.2 Proprietary IP & Algorithmic Engines Valuation (EGP 6,000,000)")
    md_lines.append("### 5.2 Proprietary IP & Algorithmic Engines Valuation (EGP 6,000,000)\n")

    ip_val_p1 = (
        "In addition to direct replacement labor, an independent valuation of MJRH's 6 proprietary intellectual property engines establishes "
        "an intangible asset value of EGP 6,000,000. This valuation is derived from the commercial cost savings, revenue leakage prevention, "
        "and competitive differentiation these algorithmic engines provide to enterprise laundry plants and future horizontal licensees:"
    )
    add_body_paragraph(doc, ip_val_p1, bold_prefix="EGP 6.00 Million Algorithmic Asset Base: ")
    md_lines.append("**EGP 6.00 Million Algorithmic Asset Base:** " + ip_val_p1 + "\n")

    # --- TABLE: IP VALUATION ---
    table_ip = doc.add_table(rows=7, cols=3)
    table_ip.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_ip.autofit = False
    col_w_ip = [Inches(2.2), Inches(3.1), Inches(1.5)]
    
    style_table_header(table_ip.rows[0], col_w_ip, ["Proprietary Algorithmic IP Engine", "Commercial Defensibility & Revenue Impact", "Appraised Value (EGP)"])
    
    ip_data = [
        ("1. APDO Operational Governance Engine", "Guarantees 100% audit compliance and prevents operational execution without data integrity.", "1,200,000"),
        ("2. Touch Hyper-Automation Engine", "Eliminates barcode scanner hardware costs and accelerates workstation throughput by 45%.", "1,100,000"),
        ("3. Dynamic Surge Scheduling Engine", "Prevents plant overload, eliminates SLA delivery breaches, and captures express surcharge revenue.", "1,000,000"),
        ("4. Double-Entry Accounting Ledger", "Automates tax compliance, eliminates cashier cash leakage, and manages digital tip liabilities.", "1,000,000"),
        ("5. Sovereign RBAC & Hybrid Governance", "Enforces German 10-department security, absolute owner sovereignty, and multi-tenant isolation.", "900,000"),
        ("6. Autonomous Telemetry & Self-Healing", "Reduces plant IT support costs by 80% via automated exception logging and error sanitization.", "800,000")
    ]
    
    for idx, (eng, def_imp, val) in enumerate(ip_data):
        style_table_row(table_ip.rows[idx+1], col_w_ip, [eng, def_imp, val], is_even=(idx%2==1), is_bold=False, align_right_start=2)
    
    md_lines.append("| Proprietary Algorithmic IP Engine | Commercial Defensibility & Revenue Impact | Appraised Value (EGP) |")
    md_lines.append("| :--- | :--- | :---: |")
    for eng, def_imp, val in ip_data:
        md_lines.append(f"| **{eng}** | {def_imp} | **{val}** |")
    md_lines.append("| **TOTAL IP VALUATION** | **Cumulative Appraised Intangible Asset Value** | **EGP 6,000,000** |\n")

    add_heading_2(doc, "5.3 Total Asset Coverage & Pre-Seed Valuation Justification")
    md_lines.append("### 5.3 Total Asset Coverage & Pre-Seed Valuation Justification\n")

    tot_p1 = (
        "Summing the Engineering Replacement Cost (EGP 5,610,000) and the Proprietary IP Valuation (EGP 6,000,000) yields a verified Total Asset "
        "Value of exactly EGP 11,610,000. When evaluated against the Company's Pre-Seed Pre-Money Valuation of EGP 15,000,000, this asset base "
        "establishes an institutional Asset Coverage Ratio of 77.4%. In standard enterprise software venture capital, seed-stage valuations "
        "are typically backed by less than 20% in tangible code assets, relying heavily on speculative growth multiples. MJRH's 77.4% asset "
        "floor provides institutional investors with an extraordinary downside protection mechanism, ensuring that the capital injection is "
        "secured by robust, highly developed industrial software capital."
    )
    add_body_paragraph(doc, tot_p1, bold_prefix="77.4% Institutional Asset Coverage Floor: ")
    md_lines.append("**77.4% Institutional Asset Coverage Floor:** " + tot_p1 + "\n")

    add_heading_2(doc, "5.4 Production Verification Case Study: VIP Order #ORD-2026-995")
    md_lines.append("### 5.4 Production Verification Case Study: VIP Order #ORD-2026-995\n")

    case_p1 = (
        "To empirically demonstrate that these valued architectural engines function in real-time commercial production without simulated "
        "mocks, due diligence auditors can inspect verified live sales order #ORD-2026-995, processed for VIP customer Dr. Sherif Al-Alfy "
        "(Villa 42, Choueifat Zone, South 90th Street, 5th Settlement, New Cairo). The order comprises 3 complex items: a Men's 2-Piece Suit "
        "(EGP 250), a Bridal Gown (EGP 600, quarantined for oil stain solvent remediation in cleaning), and a Silk Carpet (EGP 350), totaling "
        "an invoice value of EGP 1,200 under a 4-hour expedited SLA delivery priority."
    )
    add_body_paragraph(doc, case_p1, bold_prefix="Live Production Case Study in Fifth Settlement: ")
    md_lines.append("**Live Production Case Study in Fifth Settlement:** " + case_p1 + "\n")

    case_p2 = (
        "When Dr. Al-Alfy executed an InstaPay digital transfer of EGP 1,250, the Double-Entry Accounting Ledger and Tip Separation Engine "
        "instantly triggered across the Supabase kernel: exactly EGP 1,200 was credited to recognized commercial laundry revenue, while the "
        "remaining EGP 50 was automatically separated and routed as a driver tip liability owed to assigned courier Mahmoud Said. Throughout "
        "processing, the station item queries in `cleaning.tsx` and `ironing.tsx` strictly enforced `.in(\"service_type\", [\"cleaning\", \"ironing\", "
        "\"both\"])`, guaranteeing 100% item uniformity across workstations without splitting the VIP order. This real-world execution "
        "proves that MJRH's valued IP engines operate with surgical precision in production."
    )
    add_body_paragraph(doc, case_p2, bold_prefix="Algorithmic Accounting & Workflow Uniformity in Action: ")
    md_lines.append("**Algorithmic Accounting & Workflow Uniformity in Action:** " + case_p2 + "\n")

    # =========================================================================
    # SECTION 6: PRE-SEED INVESTMENT MEMORANDUM & DEAL STRUCTURE
    # =========================================================================
    add_heading_1(doc, "SECTION 6: PRE-SEED INVESTMENT MEMORANDUM & DEAL STRUCTURE")
    md_lines.append("## SECTION 6: PRE-SEED INVESTMENT MEMORANDUM & DEAL STRUCTURE\n")

    inv_p1 = (
        "MJRH INDUSTRIAL REVOLUTION is offering institutional venture capital funds, strategic private equity partners, and qualified family "
        "offices the opportunity to participate in its Pre-Seed institutional financing round. The transaction is engineered with strict "
        "financial discipline, clean equity mechanics, and absolute alignment between founder sovereignty and investor institutional protection."
    )
    add_body_paragraph(doc, inv_p1, bold_prefix="Institutional Pre-Seed Financing Overview: ")
    md_lines.append("**Institutional Pre-Seed Financing Overview:** " + inv_p1 + "\n")

    add_heading_2(doc, "6.1 Core Transaction Parameters")
    md_lines.append("### 6.1 Core Transaction Parameters\n")

    trans_p1 = (
        "The Company seeks to raise a total of EGP 3,000,000 in Pre-Seed growth capital. The round is priced at a fixed Pre-Money Enterprise "
        "Valuation of EGP 15,000,000, resulting in an EGP 18,000,000 Post-Money Valuation. Under this structure, incoming institutional investors "
        "will acquire exactly 16.67% of the newly issued preferred equity of the Company. This valuation is directly supported by the EGP "
        "11.61 Million asset floor established in Section 5, ensuring that investors acquire equity at a 1.29x multiple of verified existing "
        "engineering and IP capital."
    )
    add_body_paragraph(doc, trans_p1, bold_prefix="Valuation & Equity Structure: ")
    md_lines.append("**Valuation & Equity Structure:** " + trans_p1 + "\n")

    # --- TABLE: DEAL PARAMETERS ---
    table_deal = doc.add_table(rows=6, cols=3)
    table_deal.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_deal.autofit = False
    col_w_deal = [Inches(2.2), Inches(2.3), Inches(2.3)]
    
    style_table_header(table_deal.rows[0], col_w_deal, ["Investment Parameter", "Financial & Equity Value", "Institutional Due Diligence Note"])
    
    deal_data = [
        ("Funding Round Stage", "Pre-Seed Institutional Financing", "First institutional equity capital injection."),
        ("Target Capital Raise", "EGP 3,000,000 (~$62,500 US)", "Dedicated growth capital; zero debt leverage."),
        ("Pre-Money Enterprise Valuation", "EGP 15,000,000", "Backed by EGP 11.61M in verified code & IP assets."),
        ("Post-Money Enterprise Valuation", "EGP 18,000,000", "Establishes institutional post-money benchmark."),
        ("Investor Equity Dilution", "16.67% Preferred Equity", "Clean dilution preserving 83.33% founder incentives.")
    ]
    
    for idx, (p, v, n) in enumerate(deal_data):
        style_table_row(table_deal.rows[idx+1], col_w_deal, [p, v, n], is_even=(idx%2==1), is_bold=(idx==1 or idx==2 or idx==4), align_right_start=2)
    
    md_lines.append("| Investment Parameter | Financial & Equity Value | Institutional Due Diligence Note |")
    md_lines.append("| :--- | :---: | :--- |")
    for p, v, n in deal_data:
        md_lines.append(f"| **{p}** | {v} | {n} |")
    md_lines.append("\n---\n")

    add_heading_2(doc, "6.2 Use of Funds Allocation Schedule")
    md_lines.append("### 6.2 Use of Funds Allocation Schedule\n")

    use_p1 = (
        "To ensure maximum capital efficiency and accelerate commercial ARR scaling, the EGP 3,000,000 Pre-Seed proceeds will be deployed across "
        "four strictly governed operational pillars over an 18-month execution runway: Engineering & R&D (45% | EGP 1,350,000) dedicated to "
        "expanding AI Advisor telemetry, horizontal vertical adaptations, and automated testing suites; Go-To-Market & B2B Sales (30% | EGP "
        "900,000) deployed into direct enterprise B2B sales teams, commercial onboarding incentives, and performance marketing; Cloud "
        "Infrastructure & Security (15% | EGP 450,000) allocated to Supabase enterprise database scaling, Vercel multi-region CDN clustering, "
        "and ISO/GDPR security audits; Working Capital & Operational Reserve (10% | EGP 300,000) maintained as a liquidity buffer."
    )
    add_body_paragraph(doc, use_p1, bold_prefix="Strict Institutional Capital Allocation: ")
    md_lines.append("**Strict Institutional Capital Allocation:** " + use_p1 + "\n")

    # --- TABLE: USE OF FUNDS ---
    table_use = doc.add_table(rows=6, cols=4)
    table_use.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_use.autofit = False
    col_w_use = [Inches(2.2), Inches(1.1), Inches(1.3), Inches(2.2)]
    
    style_table_header(table_use.rows[0], col_w_use, ["Capital Deployment Pillar", "Allocation (%)", "Amount (EGP)", "Strategic Operational Milestone"])
    
    use_data = [
        ("Engineering & Product R&D", "45.0%", "1,350,000", "Scale 6 IP engines & deploy horizontal industrial adaptations."),
        ("Go-To-Market & B2B Sales", "30.0%", "900,000", "Onboard 100+ commercial laundry plants & achieve EGP 3.6M ARR."),
        ("Cloud Infrastructure & Security", "15.0%", "450,000", "Enterprise Supabase scaling, multi-region CDN, & SOC2/ISO audits."),
        ("Working Capital & Reserve", "10.0%", "300,000", "Maintain unencumbered operational liquidity buffer."),
        ("TOTAL CAPITAL RAISE", "100.0%", "3,000,000", "Fully funds 18-month runway to institutional Series A readiness.")
    ]
    
    for idx, (p, a, amt, m) in enumerate(use_data):
        style_table_row(table_use.rows[idx+1], col_w_use, [p, a, amt, m], is_even=(idx%2==1), is_bold=(idx==4), align_right_start=1)
    
    md_lines.append("| Capital Deployment Pillar | Allocation (%) | Amount (EGP) | Strategic Operational Milestone |")
    md_lines.append("| :--- | :---: | :---: | :--- |")
    for p, a, amt, m in use_data:
        md_lines.append(f"| **{p}** | {a} | **{amt}** | {m} |")
    md_lines.append("\n---\n")

    add_heading_2(doc, "6.3 Operational Runway & Series A Readiness Milestones")
    md_lines.append("### 6.3 Operational Runway & Series A Readiness Milestones\n")

    run_p1 = (
        "The EGP 3,000,000 capital injection guarantees MJRH an unencumbered 18-month operational runway without requiring supplementary equity "
        "or debt financing. By Month 18, the Company targets reaching EGP 3,600,000 in Annual Recurring Revenue (ARR) across 100+ multi-branch "
        "commercial laundry enterprises. Achieving these metrics establishes definitive institutional Series A readiness, positioning MJRH for "
        "a targeted EGP 30,000,000 Series A financing round at an estimated EGP 150,000,000 Pre-Money Valuation—delivering a projected 10x "
        "valuation markup for Pre-Seed institutional participants."
    )
    add_body_paragraph(doc, run_p1, bold_prefix="18-Month Runway to Series A Valuation Markup: ")
    md_lines.append("**18-Month Runway to Series A Valuation Markup:** " + run_p1 + "\n")

    add_heading_2(doc, "6.4 Cap Table Governance & Institutional Protective Provisions")
    md_lines.append("### 6.4 Cap Table Governance & Institutional Protective Provisions\n")

    cap_p1 = (
        "Post-transaction cap table governance is structured to preserve founder entrepreneurial velocity while embedding standard institutional "
        "protective provisions. Following the 16.67% Pre-Seed issuance, Founder & CEO Mohamed Riyad retains 83.33% common equity ownership, "
        "subject to a standard 4-year founder vesting schedule with a 1-year cliff. Institutional preferred equity holders will receive "
        "customary protective rights, including: (1) A 1x non-participating liquidation preference; (2) Pro-rata pre-emptive rights to participate "
        "in subsequent equity financings; (3) Broad information rights requiring monthly financial statements and quarterly executive QBR reports; "
        "and (4) Customary consent rights over corporate liquidations, debt incurrence exceeding EGP 500,000, or alterations to share class rights."
    )
    add_body_paragraph(doc, cap_p1, bold_prefix="Institutional Governance & Protective Provisions: ")
    md_lines.append("**Institutional Governance & Protective Provisions:** " + cap_p1 + "\n")

print("Sections 04 to 06 generated (Expanded).")
