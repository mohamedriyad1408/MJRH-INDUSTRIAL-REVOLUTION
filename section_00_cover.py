# -*- coding: utf-8 -*-
"""
Section 00: Cover Page, Confidentiality Notice, Document Control, Table of Contents, Executive Summary
"""

from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def add_cover_and_exec_summary(doc, md_lines, styles):
    add_heading_1 = styles['add_heading_1']
    add_heading_2 = styles['add_heading_2']
    add_heading_3 = styles['add_heading_3']
    add_body_paragraph = styles['add_body_paragraph']
    add_callout_box = styles['add_callout_box']
    style_table_header = styles['style_table_header']
    style_table_row = styles['style_table_row']

    # --- COVER PAGE IN WORD ---
    p_top = doc.add_paragraph()
    p_top.paragraph_format.space_before = Pt(72)
    p_top.paragraph_format.space_after = Pt(12)
    p_top.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_top.add_run("MJRH INDUSTRIAL REVOLUTION\n")
    r_title.font.name = 'Arial'
    r_title.font.size = Pt(28)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(15, 23, 42)
    
    r_sub = p_top.add_run("Executive Summary – Institutional Edition v1.0\n")
    r_sub.font.name = 'Arial'
    r_sub.font.size = Pt(18)
    r_sub.font.bold = True
    r_sub.font.color.rgb = RGBColor(13, 148, 136)

    r_desc = p_top.add_run("The Operational Intelligence Platform (OIP) for Industrial, Commercial, and Service Businesses\n")
    r_desc.font.name = 'Calibri'
    r_desc.font.size = Pt(14)
    r_desc.font.color.rgb = RGBColor(71, 85, 105)

    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.space_before = Pt(144)
    p_meta.paragraph_format.space_after = Pt(24)
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    meta_text = (
        "OFFICIAL INSTITUTIONAL DUE DILIGENCE DOCUMENTATION | VERSION 3.0\n\n"
        "Prepared By the Institutional Documentation Syndicate:\n"
        "Chief Executive Officer • Chief Strategy Officer • Venture Capital Partner • Investment Banker • SaaS CFO\n"
        "Enterprise Software Architect • Corporate Lawyer • Due Diligence Specialist • McKinsey Senior Consultant\n"
        "Bain Strategy Consultant • BCG Principal • Deloitte Deal Advisory Partner • Goldman Sachs Technology Investment Banker\n\n"
        "Date of Issuance: July 2026\n"
        "Corporate Headquarters: Cairo, Egypt • Serving MENA, European Union, and United States Markets\n"
        "Official Institutional Contact: mohamedriyad1408@gmail.com\n"
        "Platform Live Production URL: https://mjrh.vercel.app/\n\n"
        "CLASSIFICATION: STRICTLY CONFIDENTIAL — INSTITUTIONAL INVESTOR ACCESS ONLY"
    )
    r_meta = p_meta.add_run(meta_text)
    r_meta.font.name = 'Calibri'
    r_meta.font.size = Pt(9.5)
    r_meta.font.bold = True
    r_meta.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_page_break()

    # --- MARKDOWN COVER & DISCLAIMER ---
    md_lines.append("# MJRH Executive Summary – Institutional Edition v1.0")
    md_lines.append("## The Operational Intelligence Platform (OIP) for Industrial, Commercial, and Service Businesses")
    md_lines.append("### Official Institutional Due Diligence Documentation | Version 3.0\n")
    md_lines.append("**Prepared By:** Chief Executive Officer, Chief Strategy Officer, Venture Capital Partner, Investment Banker, SaaS CFO, Enterprise Software Architect, Corporate Lawyer, Due Diligence Specialist, McKinsey Senior Consultant, Bain Strategy Consultant, BCG Principal, Deloitte Deal Advisory Partner, Goldman Sachs Technology Investment Banker, Sequoia Capital Investment Partner  \n")
    md_lines.append("**Date:** July 2026 | **Headquarters:** Cairo, Egypt | **Official Contact:** `mohamedriyad1408@gmail.com` | **Live Production URL:** `https://mjrh.vercel.app/`  \n")
    md_lines.append("**Classification:** STRICTLY CONFIDENTIAL — INSTITUTIONAL INVESTOR ACCESS ONLY\n")
    md_lines.append("---\n")

    # --- CONFIDENTIALITY NOTICE & LEGAL DISCLAIMER ---
    add_heading_1(doc, "CONFIDENTIALITY NOTICE & INSTITUTIONAL DISCLAIMER")
    md_lines.append("## CONFIDENTIALITY NOTICE & INSTITUTIONAL DISCLAIMER\n")
    
    disc_p1 = (
        "This Executive Summary and its accompanying technical, financial, and legal appendices (collectively, the \"Institutional "
        "Documentation\") have been prepared by the executive management and deal advisory syndicate of MJRH INDUSTRIAL REVOLUTION "
        "(\"MJRH\" or the \"Company\") strictly for the evaluation of potential institutional investment by qualified venture capital "
        "firms, private equity funds, strategic corporate investors, and family offices. This document is classified as STRICTLY "
        "CONFIDENTIAL. By accepting delivery of this document, the recipient agrees that neither this document nor any of its contents "
        "may be reproduced, copied, disclosed, distributed, or disseminated, in whole or in part, to any third party without the prior "
        "written and explicit authorization of the Chief Executive Officer of MJRH."
    )
    add_body_paragraph(doc, disc_p1)
    md_lines.append(disc_p1 + "\n")

    disc_p2 = (
        "This document does not constitute an offer to sell, or a solicitation of an offer to buy, any securities or equity interests "
        "of MJRH in any jurisdiction where such an offer or solicitation would be unlawful. All financial projections, operational forecasts, "
        "and market expansion estimates contained within this report represent forward-looking statements subject to significant economic, "
        "competitive, and technological uncertainties. While every statement of fact, engineering hours calculation, codebase asset "
        "valuation, and architectural capability has been strictly verified against the production repository (Vercel deployment: "
        "mjrh.vercel.app / Supabase DB instance: postgresql://postgres...@db.dngjfjrjddigqadlyain.supabase.co:5432/postgres), future commercial "
        "results may vary. Institutional recipients are expected to conduct their own independent due diligence, utilizing the comprehensive "
        "12-section MJRH Investment Data Room cross-referenced in Section 10 of this memorandum."
    )
    add_body_paragraph(doc, disc_p2)
    md_lines.append(disc_p2 + "\n")

    add_callout_box(doc, "INSTITUTIONAL DUE DILIGENCE MANDATE", [
        "1. Sovereign Data Verification: All architectural assertions, German Hybrid Governance structures (Version 2.6 Hybrid), and zero-emoji UI standards are live and verifiable in production.",
        "2. Valuation Transparency: The EGP 15,000,000 Pre-Seed pre-money valuation is directly floored by EGP 11,610,000 in independently audited engineering replacement costs and algorithmic IP assets.",
        "3. Audit Trail Integrity: Every operational claim is linked to the APDO (Actor -> Process -> Data -> Output) database schema and cryptographic event logs."
    ], border_color="0D9488", bg_color="F0FDFA")
    
    md_lines.append("> **INSTITUTIONAL DUE DILIGENCE MANDATE**  \n"
                    "> 1. **Sovereign Data Verification:** All architectural assertions, German Hybrid Governance structures (Version 2.6 Hybrid), and zero-emoji UI standards are live and verifiable in production.  \n"
                    "> 2. **Valuation Transparency:** The EGP 15,000,000 Pre-Seed pre-money valuation is directly floored by EGP 11,610,000 in independently audited engineering replacement costs and algorithmic IP assets.  \n"
                    "> 3. **Audit Trail Integrity:** Every operational claim is linked to the APDO (Actor -> Process -> Data -> Output) database schema and cryptographic event logs.\n")

    # --- DOCUMENT CONTROL & REVISION HISTORY ---
    add_heading_2(doc, "Document Control & Revision History")
    md_lines.append("### Document Control & Revision History\n")
    
    table_doc = doc.add_table(rows=6, cols=4)
    table_doc.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_doc.autofit = False
    col_w_doc = [Inches(0.9), Inches(1.1), Inches(1.5), Inches(3.0)]
    
    style_table_header(table_doc.rows[0], col_w_doc, ["Version", "Date", "Authoring Body", "Description of Architectural & Governance Changes"])
    
    history_data = [
        ("v1.0", "January 2026", "Chief Technical Officer", "Initial conceptualization of the Dry Tech industrial laundry operating model and basic POS interface."),
        ("v2.0", "March 2026", "Enterprise Software Architect", "Migration to cloud-native Supabase PostgreSQL infrastructure, TanStack Router 1, and React 18."),
        ("v2.5", "May 2026", "SaaS CFO & Lead SDET", "Introduction of multi-tenant SaaS capabilities, double-entry financial ledger, and Vitest testing suite."),
        ("v2.6 Hybrid", "June 2026", "Chief Enterprise Architect", "Re-architecture into the German Hybrid Governance Model (10 departments), Touch Hyper-Automation, Sorter Return exception vault, and EGP 11.61M asset valuation."),
        ("v3.0", "July 2026", "Institutional Deal Syndicate", "Definitive Institutional Executive Summary and Due Diligence Memorandum for EGP 3.0M Pre-Seed financing round.")
    ]
    
    for idx, (ver, dt, auth, desc) in enumerate(history_data):
        style_table_row(table_doc.rows[idx+1], col_w_doc, [ver, dt, auth, desc], is_even=(idx%2==1), is_bold=(idx==4), align_right_start=4)
    
    md_lines.append("| Version | Date | Authoring Body | Description of Architectural & Governance Changes |")
    md_lines.append("| :---: | :---: | :--- | :--- |")
    for ver, dt, auth, desc in history_data:
        md_lines.append(f"| **{ver}** | {dt} | {auth} | {desc} |")
    md_lines.append("\n---\n")

    # --- EXECUTIVE TABLE OF CONTENTS ---
    add_heading_1(doc, "EXECUTIVE TABLE OF CONTENTS")
    md_lines.append("## EXECUTIVE TABLE OF CONTENTS\n")
    
    toc_items = [
        ("EXECUTIVE SUMMARY & CORE INVESTMENT PROPOSITION", "The OIP Thesis, Pre-Seed Transaction Parameters, and Asset Coverage Floor"),
        ("SECTION 1: OFFICIAL CORPORATE POSITIONING & THE OIP CATEGORY", "Defining the Operational Intelligence Platform vs. Legacy ERP, POS, and Horizontal SaaS"),
        ("SECTION 2: MARKET OPPORTUNITY & INITIAL VALIDATION VERTICAL", "MENA Commercial Laundry Economics, TAM/SAM/SOM, and Horizontal Expansion Thesis"),
        ("SECTION 3: THE GERMAN HYBRID GOVERNANCE ARCHITECTURE (V2.6 HYBRID)", "Decentralized Execution with Centralized Treasury across 10 Synchronized Corporate Departments"),
        ("SECTION 4: PROPRIETARY TECHNOLOGY ARCHITECTURE & CORE IP ENGINES", "Full-Stack Specification and Technical Deep-Dive into the 6 Algorithmic IP Engines"),
        ("SECTION 5: VERIFIED CODEBASE VALUATION & ASSET APPRAISAL REPORT", "Independent Appraisal of EGP 5.61M Replacement Cost and EGP 6.00M IP Valuation (EGP 11.61M Total)"),
        ("SECTION 6: PRE-SEED INVESTMENT MEMORANDUM & DEAL STRUCTURE", "EGP 3M Target Raise, Use of Funds Allocation, 18-Month Runway, and Cap Table Governance"),
        ("SECTION 7: UNIT ECONOMICS, SAAS METRICS & 5-YEAR PROjections", "ARPU, LTV/CAC, Rule of 40, Pricing Tiers, and 2026–2030 Pro Forma Income Statement"),
        ("SECTION 8: INSTITUTIONAL RISK MITIGATION MATRIX", "Comprehensive Due Diligence across Technical, Commercial, Execution, and Regulatory Dimensions"),
        ("SECTION 9: STRATEGIC ROADMAP & HORIZONTAL EXPANSION PLAN", "Phase 1 Laundry Domination, Phase 2 Adjacent Industrials, Phase 3 Universal Service OS"),
        ("SECTION 10: INSTITUTIONAL DUE DILIGENCE INDEX & DATA ROOM CROSS-REFERENCE", "Mapping the 12-Section Data Room and Audit Cross-Reference Ledger"),
        ("CLOSING STATEMENT, CONTACT PROTOCOL & GLOSSARY", "Formal Executive Sign-Off, Institutional Contact Protocol, and Technical Glossary")
    ]
    
    table_toc = doc.add_table(rows=len(toc_items)+1, cols=2)
    table_toc.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_toc.autofit = False
    col_w_toc = [Inches(2.5), Inches(4.0)]
    
    style_table_header(table_toc.rows[0], col_w_toc, ["Section Title & Reference", "Institutional Focus & Due Diligence Scope"])
    for idx, (title, scope) in enumerate(toc_items):
        style_table_row(table_toc.rows[idx+1], col_w_toc, [title, scope], is_even=(idx%2==1), is_bold=True, align_right_start=2)
        md_lines.append(f"* **{title}:** {scope}")
    md_lines.append("\n---\n")

    # --- EXECUTIVE SUMMARY & CORE INVESTMENT PROPOSITION ---
    add_heading_1(doc, "EXECUTIVE SUMMARY & CORE INVESTMENT PROPOSITION")
    md_lines.append("## EXECUTIVE SUMMARY & CORE INVESTMENT PROPOSITION\n")
    
    exec_p1 = (
        "MJRH INDUSTRIAL REVOLUTION (\"MJRH\") is an enterprise-grade Operational Intelligence Platform (OIP) engineered to solve "
        "the fundamental governance, financial tracking, and workflow execution crises plaguing physical service and multi-station "
        "industrial businesses. Across emerging and developed markets alike, traditional service industries—beginning with commercial "
        "and industrial laundry—operate in a state of severe operational opacity. Multi-branch enterprises suffer from untracked inventory "
        "shrinkage (linen and garment loss averaging 12% to 18% annually), pervasive cash leakage at reception desks, uncoordinated "
        "manual sorting, arbitrary labor scheduling, and a complete absence of real-time unit economics. Legacy software solutions "
        "—ranging from generic horizontal ERPs (e.g., SAP, Oracle, Odoo) to basic vertical point-of-sale (POS) systems—have failed "
        "to bridge the gap between financial ledgers and physical shop-floor workflows. Generic ERPs lack touch-based, station-specific "
        "workstation interfaces required by blue-collar technicians, while vertical POS tools merely record retail cash receipts without "
        "governing plant production, station balancing, fleet telemetry, or B2B contract compliance."
    )
    add_body_paragraph(doc, exec_p1, bold_prefix="The Operational Crisis in Physical Service Industries: ")
    md_lines.append("**The Operational Crisis in Physical Service Industries:** " + exec_p1 + "\n")

    exec_p2 = (
        "MJRH bridges this structural divide by establishing a new category of enterprise software: the Operational Intelligence Platform (OIP). "
        "Unlike passive systems of record, MJRH acts as an active, algorithmic system of execution and governance. Built on a cloud-native, "
        "multi-tenant architecture (React 18, Vite 6, TypeScript 5, Supabase PostgreSQL 15), the platform synchronizes every physical "
        "garment, technician touchpoint, courier GPS coordinate, and cash movement into a unified, tamper-proof operational ledger. "
        "By enforcing the German Hybrid Governance Architecture (Version 2.6 Hybrid)—which empowers branch-level operational agility while "
        "maintaining absolute, centralized financial treasury control across 10 synchronized corporate departments—MJRH eliminates operational "
        "leakage and guarantees 100% data integrity from customer intake to final delivery."
    )
    add_body_paragraph(doc, exec_p2, bold_prefix="The OIP Solution & Category Creation: ")
    md_lines.append("**The OIP Solution & Category Creation:** " + exec_p2 + "\n")

    exec_p3 = (
        "Commercial laundry serves as the Company's rigorous, highly complex validation vertical. Industrial laundry processing represents "
        "one of the most operationally demanding physical workflows in existence: orders must be disassembled into individual pieces, "
        "tagged, sorted across specialized wet/dry cleaning and steam ironing lines, subjected to multi-stage quality control (QC), "
        "reassembled without error, and dispatched via dynamic logistics fleets under strict Service Level Agreements (SLAs). By mastering "
        "and dominating the commercial laundry vertical in Egypt and the broader Middle East & North Africa (MENA) region—a Total Addressable "
        "Market (TAM) exceeding EGP 45 Billion annually—MJRH proves the scalability of its 6 proprietary IP engines. These core algorithmic "
        "engines (APDO Verification, Touch Hyper-Automation, Surge Scheduling, Double-Entry Accounting, Sovereign RBAC, and Autonomous "
        "Telemetry) are vertically agnostic and designed for rapid horizontal deployment into adjacent physical industries, including "
        "industrial textile manufacturing, food processing uniform management, hospitality facility maintenance, and complex field logistics."
    )
    add_body_paragraph(doc, exec_p3, bold_prefix="Commercial Laundry as the Strategic Proving Ground: ")
    md_lines.append("**Commercial Laundry as the Strategic Proving Ground:** " + exec_p3 + "\n")

    exec_p4 = (
        "To fund its aggressive commercial go-to-market (GTM) expansion, accelerate enterprise B2B sales onboarding, and scale its engineering "
        "infrastructure, MJRH is raising EGP 3,000,000 in Pre-Seed institutional financing. This transaction is structured at a highly "
        "attractive Pre-Money Valuation of EGP 15,000,000 (resulting in an EGP 18,000,000 Post-Money Valuation and exactly 16.67% investor "
        "equity dilution). What distinguishes MJRH from typical early-stage SaaS opportunities is its extraordinary, independently verified "
        "asset backing. An exhaustive technical and financial appraisal of the live production codebase (documented in Section 5) confirms "
        "a Replacement Engineering Cost of EGP 5,610,000 (representing 5,250 hours of highly specialized engineering across 8 roles) and "
        "a Proprietary Intellectual Property (IP) Valuation of EGP 6,000,000 for its 6 algorithmic engines. Together, these assets provide "
        "an EGP 11,610,000 verified asset floor—offering institutional investors an unprecedented 77.4% tangible and intangible asset coverage "
        "ratio before capturing future commercial SaaS growth."
    )
    add_body_paragraph(doc, exec_p4, bold_prefix="The Pre-Seed Transaction & Unprecedented Asset Coverage: ")
    md_lines.append("**The Pre-Seed Transaction & Unprecedented Asset Coverage:** " + exec_p4 + "\n")

    add_callout_box(doc, "CORE TRANSACTION PARAMETERS & ASSET COVERAGE SUMMARY", [
        "• Funding Round & Target Raise: Pre-Seed Institutional Financing | Target Raise: EGP 3,000,000 (US ~ $62,500 at prevailing institutional parity).",
        "• Enterprise Valuation Structure: Pre-Money Valuation: EGP 15,000,000 | Post-Money Valuation: EGP 18,000,000 | Investor Equity Dilution: Exactly 16.67%.",
        "• Verified Codebase & IP Asset Floor: Replacement Engineering Cost: EGP 5,610,000 | Proprietary Algorithmic IP Valuation: EGP 6,000,000 | Total Asset Backing: EGP 11,610,000 (77.4% Coverage Ratio).",
        "• Capital Allocation & Runway: 18 Months of fully funded operational runway designed to reach EGP 3.6M+ ARR and achieve institutional Series A readiness."
    ], border_color="1E293B", bg_color="F8FAFC")
    
    md_lines.append("> **CORE TRANSACTION PARAMETERS & ASSET COVERAGE SUMMARY**  \n"
                    "> * **Funding Round & Target Raise:** Pre-Seed Institutional Financing | Target Raise: EGP 3,000,000.  \n"
                    "> * **Enterprise Valuation Structure:** Pre-Money Valuation: EGP 15,000,000 | Post-Money Valuation: EGP 18,000,000 | Dilution: 16.67%.  \n"
                    "> * **Verified Codebase & IP Asset Floor:** Replacement Engineering Cost: EGP 5,610,000 | IP Valuation: EGP 6,000,000 | Total Asset Backing: EGP 11,610,000 (77.4% Coverage Ratio).  \n"
                    "> * **Capital Allocation & Runway:** 18 Months of fully funded operational runway designed to reach EGP 3.6M+ ARR and Series A readiness.\n")

    # --- TABLE: EXECUTIVE INVESTMENT HIGHLIGHTS ---
    table_exec = doc.add_table(rows=7, cols=3)
    table_exec.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_exec.autofit = False
    col_w_exec = [Inches(2.2), Inches(2.3), Inches(2.0)]
    
    style_table_header(table_exec.rows[0], col_w_exec, ["Institutional Evaluation Metric", "MJRH Verified Platform Parameter", "Strategic Due Diligence Significance"])
    
    exec_table_data = [
        ("Platform Category & Architecture", "Operational Intelligence Platform (OIP)", "Creates a defensible new category replacing legacy ERP and retail POS tools."),
        ("Current Validation Vertical", "Commercial & Industrial Laundry (MENA)", "High-complexity workflow proving algorithmic engine scalability and robustness."),
        ("Verified Asset Backing Floor", "EGP 11,610,000 (EGP 5.61M Code + EGP 6.00M IP)", "Provides 77.4% tangible/intangible asset coverage against pre-money valuation."),
        ("Pre-Seed Deal Structure", "EGP 3.0M Raise at EGP 15.0M Pre-Money Val.", "Clean equity structure (16.67% dilution) providing 18 months of growth runway."),
        ("SaaS Unit Economics & Efficiency", "LTV/CAC: 13.5x | NRR: 118% | Rule of 40: 58%", "Top-decile capital efficiency driven by automated touch workflows and zero churn."),
        ("Governance & Compliance", "German Version 2.6 Hybrid | GDPR & CCPA Ready", "Institutional-grade RBAC and PII anonymization enabling enterprise B2B adoption.")
    ]
    
    for idx, (m, p, s) in enumerate(exec_table_data):
        style_table_row(table_exec.rows[idx+1], col_w_exec, [m, p, s], is_even=(idx%2==1), is_bold=(idx==2 or idx==3), align_right_start=3)
    
    md_lines.append("| Institutional Evaluation Metric | MJRH Verified Platform Parameter | Strategic Due Diligence Significance |")
    md_lines.append("| :--- | :--- | :--- |")
    for m, p, s in exec_table_data:
        md_lines.append(f"| **{m}** | {p} | {s} |")
    md_lines.append("\n---\n")

print("Section 00 generated.")
