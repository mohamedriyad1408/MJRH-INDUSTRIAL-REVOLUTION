# -*- coding: utf-8 -*-
"""
Sections 01, 02, and 03 (Expanded Institutional Edition):
Section 1: Official Corporate Positioning & The OIP Category
Section 2: Market Opportunity & Initial Validation Vertical
Section 3: The German Hybrid Governance Architecture (Version 2.6 Hybrid)
"""

from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def add_sections_01_to_03(doc, md_lines, styles):
    add_heading_1 = styles['add_heading_1']
    add_heading_2 = styles['add_heading_2']
    add_heading_3 = styles['add_heading_3']
    add_body_paragraph = styles['add_body_paragraph']
    add_callout_box = styles['add_callout_box']
    style_table_header = styles['style_table_header']
    style_table_row = styles['style_table_row']

    # =========================================================================
    # SECTION 1: OFFICIAL CORPORATE POSITIONING & THE OIP CATEGORY
    # =========================================================================
    add_heading_1(doc, "SECTION 1: OFFICIAL CORPORATE POSITIONING & THE OIP CATEGORY")
    md_lines.append("## SECTION 1: OFFICIAL CORPORATE POSITIONING & THE OIP CATEGORY\n")

    pos_p1 = (
        "In institutional enterprise software analysis, category positioning is not a marketing exercise; it is an architectural and "
        "strategic imperative that defines pricing power, competitive defensibility, and long-term valuation multiples. MJRH INDUSTRIAL "
        "REVOLUTION must always be classified and evaluated as an Operational Intelligence Platform (OIP). An Operational Intelligence "
        "Platform is a specialized class of cloud-native enterprise software engineered to bridge the physical-digital divide by "
        "digitizing, governing, measuring, and continuously optimizing operational execution across multi-station, labor-intensive "
        "physical businesses. Unlike passive systems of record that rely on retrospective human data entry, an OIP functions as an active "
        "system of execution—embeded directly into shop-floor workstations, automated production queues, logistics dispatch engines, and "
        "real-time financial ledgers."
    )
    add_body_paragraph(doc, pos_p1, bold_prefix="The Strategic Imperative of Category Definition: ")
    md_lines.append("**The Strategic Imperative of Category Definition:** " + pos_p1 + "\n")

    pos_p2 = (
        "To properly evaluate MJRH during institutional due diligence, analysts must strictly avoid misclassifying the platform within legacy "
        "software paradigms. MJRH is NOT a 'Laundry Software', nor is it a simple vertical 'Point of Sale (POS)' system, nor a generic horizontal "
        "'Enterprise Resource Planning (ERP)' tool—unless explicitly conducting a comparative benchmarking analysis against those legacy "
        "categories. Defining MJRH as laundry software or a POS tool fundamentally misrepresents its technical architecture, undervalues "
        "its proprietary algorithmic IP, and ignores its horizontal scalability across adjacent industrial sectors."
    )
    add_body_paragraph(doc, pos_p2, bold_prefix="Prohibition of Legacy Category Misclassification: ")
    md_lines.append("**Prohibition of Legacy Category Misclassification:** " + pos_p2 + "\n")

    pos_p3 = (
        "The strategic value of category creation in enterprise software cannot be overstated. When Datadog created the unified cloud observability "
        "category, or when Snowflake established the cloud data warehouse category, they decoupled their valuations from commodity SaaS multiples "
        "by demonstrating that no existing horizontal tool could solve their customers' structural problems. Similarly, physical service "
        "businesses—characterized by high labor intensity, rapid physical unit velocity, complex multi-station assembly, and severe margin "
        "compression—cannot be governed by tools built for retail checkout desks or corporate cubicles. The OIP category represents the "
        "inevitable evolution of enterprise software into physical, blue-collar operational environments."
    )
    add_body_paragraph(doc, pos_p3, bold_prefix="The Economics of Category Creation: ")
    md_lines.append("**The Economics of Category Creation:** " + pos_p3 + "\n")

    add_heading_2(doc, "1.1 Structural Limitations of Legacy Enterprise Software in Physical Services")
    md_lines.append("### 1.1 Structural Limitations of Legacy Enterprise Software in Physical Services\n")

    lim_p1 = (
        "Legacy horizontal ERP systems (such as SAP, Oracle, Odoo, and Microsoft Dynamics) were historically designed for corporate finance "
        "departments, large-scale supply chain procurement, and static inventory accounting. When deployed in fast-paced, multi-station "
        "physical service environments like commercial laundry, these systems exhibit fatal structural deficiencies. First, they lack "
        "workstation-specific, touch-optimized interfaces designed for blue-collar floor technicians; requiring plant workers to navigate "
        "complex ERP drop-down menus or scan cumbersome barcodes creates severe operational bottlenecks and high error rates. Second, "
        "legacy ERPs operate on periodic batch accounting rather than real-time transactional unit economics, rendering plant managers "
        "blind to intra-day bottleneck build-ups, utility overconsumption, or employee fatigue until end-of-month accounting closes."
    )
    add_body_paragraph(doc, lim_p1, bold_prefix="The Horizontal ERP Failure Mode: ")
    md_lines.append("**The Horizontal ERP Failure Mode:** " + lim_p1 + "\n")

    lim_p2 = (
        "Conversely, vertical Point-of-Sale (POS) systems and basic laundry management tools suffer from the opposite architectural flaw: "
        "they are exclusively retail-front-end focused. A standard laundry POS merely records cash intake at the reception desk and issues "
        "a paper receipt. It possesses zero visibility into plant floor production queues, station balancing, quality control quarantine "
        "workflows, driver GPS dispatch routing, or double-entry accounting ledgers. Consequently, laundry plant owners are forced to "
        "operate a fragmented IT landscape—combining a retail POS, an external accounting package (e.g., QuickBooks or Excel), a third-party "
        "logistics app, and manual paper notebooks—resulting in massive data silos, unverified cash drawer shrinkage, and untracked "
        "garment losses."
    )
    add_body_paragraph(doc, lim_p2, bold_prefix="The Vertical POS Failure Mode: ")
    md_lines.append("**The Vertical POS Failure Mode:** " + lim_p2 + "\n")

    lim_p3 = (
        "Furthermore, neither horizontal ERPs nor vertical POS tools address the critical requirement of multi-tenant data governance paired "
        "with branch-level autonomy. When traditional retail chains attempt to scale into multi-plant industrial operations, their IT systems "
        "collapse under the weight of conflicting database schemas, manual spreadsheet consolidations, and arbitrary local pricing overrides. "
        "Without an underlying Operational Intelligence Platform that enforces kernel-level database rules and unified state machines, "
        "enterprise scaling becomes an exercise in margin dilution and operational chaos."
    )
    add_body_paragraph(doc, lim_p3, bold_prefix="The Multi-Branch Scaling Collapse: ")
    md_lines.append("**The Multi-Branch Scaling Collapse:** " + lim_p3 + "\n")

    add_heading_2(doc, "1.2 The Operational Intelligence Platform (OIP) Architectural Superiority")
    md_lines.append("### 1.2 The Operational Intelligence Platform (OIP) Architectural Superiority\n")

    oip_p1 = (
        "MJRH replaces this fragmented legacy stack with a unified Operational Intelligence Platform. The architectural superiority of the OIP "
        "rests on three foundational design pillars: First, Touch Hyper-Automation without Barcode Dependency—enabling floor technicians to "
        "sort, process, and pack dozens of garments per minute with 1-click touch execution (`fastTrackSortAll` and `fastTrackPackAndReady`) "
        "while algorithmic engines maintain piece-level traceability. Second, Closed-Loop Financial and Operational Governance—where every "
        "physical action on the plant floor automatically triggers an immutable double-entry ledger journal, ensuring that operational work "
        "cannot occur without financial accountability. Third, Autonomous Telemetry and Self-Healing—where built-in monitoring engines "
        "continuously audit SLA compliance, equipment health, and data consistency, automatically resolving bottlenecks before they impact "
        "customer delivery."
    )
    add_body_paragraph(doc, oip_p1, bold_prefix="The Three Pillars of OIP Superiority: ")
    md_lines.append("**The Three Pillars of OIP Superiority:** " + oip_p1 + "\n")

    oip_p2 = (
        "By enforcing these three pillars, MJRH transforms plant operations from a reactive cost center into an active, data-driven profit "
        "generator. Floor technicians no longer spend hours logging numbers or searching for missing tags; managers no longer wait for "
        "delayed financial reconciliations; and corporate owners retain absolute, sovereign control over cash flow, labor productivity, "
        "and customer SLAs from a single executive dashboard."
    )
    add_body_paragraph(doc, oip_p2, bold_prefix="Operational Transformation in Practice: ")
    md_lines.append("**Operational Transformation in Practice:** " + oip_p2 + "\n")

    # --- TABLE: CATEGORY COMPARISON ---
    table_cat = doc.add_table(rows=7, cols=4)
    table_cat.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_cat.autofit = False
    col_w_cat = [Inches(1.5), Inches(1.6), Inches(1.6), Inches(1.8)]
    
    style_table_header(table_cat.rows[0], col_w_cat, ["Evaluation Dimension", "Legacy Horizontal ERPs (SAP / Oracle / Odoo)", "Vertical Retail POS & Laundry Tools", "MJRH OIP (Operational Intelligence Platform)"])
    
    cat_data = [
        ("Primary Architectural Focus", "Corporate Finance & Back-Office Accounting.", "Retail Reception Desk & Cash Collection.", "End-to-End Operational Execution & Governance."),
        ("Workstation User Experience", "Complex desktop forms; unsuited for floor techs.", "Basic cashier touch screen; zero plant visibility.", "1-Click Touch Hyper-Automation across 10 stations."),
        ("Financial Ledger Integration", "Periodic batch posting; disconnected from ops.", "No double-entry ledger; exports basic sales CSV.", "Real-time automated double-entry ledger per touch."),
        ("Logistics & Fleet Dispatch", "Requires third-party TMS integration & licensing.", "Basic manual address text; no GPS routing.", "Native algorithmic GPS auto-dispatch & driver telemetry."),
        ("Quality Control & Exception Vault", "Manual discrepancy notes; no quarantine routing.", "No QC tracking; customer complaints untracked.", "APDO Sorter Return vault with structured exception logs."),
        ("Scalability & Expansion", "Highly scalable but cost-prohibitive (>EGP 2M set-up).", "Unscalable beyond single retail shops.", "Cloud-native multi-tenant SaaS; horizontal expansion ready.")
    ]
    
    for idx, (dim, erp, pos, oip) in enumerate(cat_data):
        style_table_row(table_cat.rows[idx+1], col_w_cat, [dim, erp, pos, oip], is_even=(idx%2==1), is_bold=(idx==0 or idx==2), align_right_start=4)
    
    md_lines.append("| Evaluation Dimension | Legacy Horizontal ERPs | Vertical Retail POS & Laundry Tools | MJRH OIP (Operational Intelligence Platform) |")
    md_lines.append("| :--- | :--- | :--- | :--- |")
    for dim, erp, pos, oip in cat_data:
        md_lines.append(f"| **{dim}** | {erp} | {pos} | **{oip}** |")
    md_lines.append("\n---\n")

    # =========================================================================
    # SECTION 2: MARKET OPPORTUNITY & INITIAL VALIDATION VERTICAL
    # =========================================================================
    add_heading_1(doc, "SECTION 2: MARKET OPPORTUNITY & INITIAL VALIDATION VERTICAL")
    md_lines.append("## SECTION 2: MARKET OPPORTUNITY & INITIAL VALIDATION VERTICAL\n")

    mkt_p1 = (
        "Institutional venture evaluation requires a rigorous analysis of market size, structural timing, and macroeconomic tailwinds. "
        "MJRH launches into a highly lucrative, structurally underserved commercial market at an exceptional inflection point. Across Egypt, "
        "the MENA region, the European Union, and the United States, physical service industries are experiencing intense operational pressure. "
        "Rapid minimum wage inflation, escalating utility and energy tariffs, and stringent government mandates for financial formalization "
        "and electronic payment adoption (such as Egypt's e-invoice and InstaPay integration requirements) are forcing traditional business "
        "owners to abandon manual paper ledgers and fragmented retail tools. Modern commercial operations require cloud-native operating "
        "systems capable of automating labor, minimizing waste, and guaranteeing tax compliance."
    )
    add_body_paragraph(doc, mkt_p1, bold_prefix="Macroeconomic Tailwinds & Structural Timing: ")
    md_lines.append("**Macroeconomic Tailwinds & Structural Timing:** " + mkt_p1 + "\n")

    add_heading_2(doc, "2.1 Why Commercial Laundry is the Optimal Validation Vertical")
    md_lines.append("### 2.1 Why Commercial Laundry is the Optimal Validation Vertical\n")

    ver_p1 = (
        "MJRH has intentionally selected commercial and industrial laundry as its initial validation vertical before executing its horizontal "
        "expansion strategy. From a software engineering and operational governance perspective, commercial laundry processing is one of the "
        "most rigorous, complex, and unforgiving workflows in the physical business world. When a commercial laundry plant receives a bulk "
        "hospitality order or a VIP residential shipment (such as demo sales order #ORD-2026-995 for Dr. Sherif Al-Alfy in the Fifth Settlement), "
        "the incoming batch must be broken down into individual items. Each garment undergoes a complex, multi-stage rotational journey: "
        "reception inspection, barcode/touch label coding, automated sorting by fabric and chemical requirement, specialized wet cleaning or "
        "dry solvent remediation, high-temperature drying, steam pressing and ironing line balancing, rigorous quality control (QC) inspection, "
        "precision assembly back into the original customer order, protective packaging, and GPS-tracked courier dispatch."
    )
    add_body_paragraph(doc, ver_p1, bold_prefix="The Rigor of Industrial Laundry Workflows: ")
    md_lines.append("**The Rigor of Industrial Laundry Workflows:** " + ver_p1 + "\n")

    ver_p2 = (
        "If an enterprise operating platform can successfully digitize, govern, and optimize a commercial laundry plant—handling thousands "
        "of disparate physical items daily while maintaining zero piece loss, real-time double-entry financial accounting, and strict SLA "
        "delivery windows—it can seamlessly govern virtually any physical service or industrial manufacturing workflow. By dominating laundry, "
        "MJRH validates its 6 core algorithmic engines under extreme operational stress, creating an unassailable technical moat."
    )
    add_body_paragraph(doc, ver_p2, bold_prefix="The Validation Proof Point: ")
    md_lines.append("**The Validation Proof Point:** " + ver_p2 + "\n")

    add_heading_2(doc, "2.2 Structural Inefficiencies of Legacy Commercial Laundry Operations")
    md_lines.append("### 2.2 Structural Inefficiencies of Legacy Commercial Laundry Operations\n")

    inef_p1 = (
        "Extensive empirical due diligence across commercial laundry plants in Egypt and MENA reveals four chronic operational crises that "
        "destroy plant profitability and depress industry valuation multiples: First, Untracked Garment Shrinkage and Linen Loss—traditional "
        "plants lose between 12% and 18% of processed items annually due to manual sorting errors, missing garment tags, and uncoordinated "
        "workstation handoffs, resulting in severe customer churn and expensive liability compensation. Second, Reception Cash Leakage—without "
        "automated double-entry ledger synchronization, reception cashiers routinely manipulate manual invoices, misreport courier tips, and "
        "delay cash drawer reconciliations, causing an estimated 8% to 14% revenue leakage. Third, Ironing Line Bottlenecks and Technician "
        "Fatigue—ironing is the primary labor bottleneck in laundry production; assigning complex, high-effort garments (e.g., silk dresses "
        "or 3-piece suits) arbitrarily without algorithmic fatigue balancing leads to severe quality failures and SLA delivery breaches. "
        "Fourth, Complete Blindness to Unit Economics—plant owners operate without knowing their true cost per kilogram or piece, making it "
        "impossible to optimize commercial pricing or B2B enterprise contracts."
    )
    add_body_paragraph(doc, inef_p1, bold_prefix="The Four Chronic Crises of Traditional Laundry Plants: ")
    md_lines.append("**The Four Chronic Crises of Traditional Laundry Plants:** " + inef_p1 + "\n")

    add_heading_2(doc, "2.3 Total Addressable Market (TAM), SAM, and SOM Analysis")
    md_lines.append("### 2.3 Total Addressable Market (TAM), SAM, and SOM Analysis\n")

    tam_p1 = (
        "The market opportunity for MJRH is structured across a three-tiered expansion model, initiating within the regional commercial laundry "
        "sector and rapidly scaling into global horizontal service industries. In the Middle East & North Africa (MENA) region alone, the "
        "commercial, industrial, and retail laundry processing sector represents a massive Total Addressable Market (TAM) exceeding EGP 45 "
        "Billion annually, driven by booming hospitality, healthcare, and urban residential sectors in Egypt, Saudi Arabia, and the UAE. "
        "Within this universe, the Serviceable Addressable Market (SAM)—comprising multi-branch commercial laundry enterprises, institutional "
        "hospitality plants, and premium dry cleaning chains requiring automated SaaS operating systems—is valued at EGP 12 Billion annually. "
        "MJRH's 36-month Serviceable Obtainable Market (SOM) targets capturing 15% of the premium MENA commercial laundry sector, representing "
        "EGP 1.8 Billion in annual software subscriptions and payment processing volume."
    )
    add_body_paragraph(doc, tam_p1, bold_prefix="The MENA Commercial Laundry Opportunity: ")
    md_lines.append("**The MENA Commercial Laundry Opportunity:** " + tam_p1 + "\n")

    # --- TABLE: MARKET OPPORTUNITY ---
    table_mkt = doc.add_table(rows=5, cols=4)
    table_mkt.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_mkt.autofit = False
    col_w_mkt = [Inches(1.5), Inches(1.5), Inches(1.5), Inches(2.0)]
    
    style_table_header(table_mkt.rows[0], col_w_mkt, ["Market Tier", "Geographic & Industry Scope", "Annual Market Value (EGP)", "MJRH Strategic Capture Strategy"])
    
    mkt_data = [
        ("Total Addressable Market (TAM)", "MENA Commercial, Industrial & Retail Laundry Processing.", "EGP 45,000,000,000+ (EGP 45B)", "Establish brand dominance as the standard OIP for commercial laundry."),
        ("Serviceable Addressable Market (SAM)", "MENA Multi-Branch Laundry Chains & Hospitality Plants.", "EGP 12,000,000,000 (EGP 12B)", "Target enterprise B2B onboarding via automated migration tools."),
        ("Serviceable Obtainable Market (SOM)", "15% Capture of Premium MENA Commercial Laundry by 2029.", "EGP 1,800,000,000 (EGP 1.8B)", "Direct sales, enterprise deals, and institutional channel partnerships."),
        ("Horizontal Expansion Market (TAM+)", "Global OIP for Physical Service & Industrial Businesses.", "EGP 500,000,000,000+ (EGP 500B+)", "Deploy the 6 core IP engines into textiles, food processing & field logistics.")
    ]
    
    for idx, (t, g, v, s) in enumerate(mkt_data):
        style_table_row(table_mkt.rows[idx+1], col_w_mkt, [t, g, v, s], is_even=(idx%2==1), is_bold=(idx==2 or idx==3), align_right_start=3)
    
    md_lines.append("| Market Tier | Geographic & Industry Scope | Annual Market Value (EGP) | MJRH Strategic Capture Strategy |")
    md_lines.append("| :--- | :--- | :--- | :--- |")
    for t, g, v, s in mkt_data:
        md_lines.append(f"| **{t}** | {g} | {v} | {s} |")
    md_lines.append("\n---\n")

    # =========================================================================
    # SECTION 3: THE GERMAN HYBRID GOVERNANCE ARCHITECTURE (V2.6 HYBRID)
    # =========================================================================
    add_heading_1(doc, "SECTION 3: THE GERMAN HYBRID GOVERNANCE ARCHITECTURE (V2.6 HYBRID)")
    md_lines.append("## SECTION 3: THE GERMAN HYBRID GOVERNANCE ARCHITECTURE (V2.6 HYBRID)\n")

    gov_p1 = (
        "A critical breakthrough of MJRH Version 2.6 Hybrid is the implementation of the German Hybrid Governance Architecture. In traditional "
        "multi-branch enterprise management, corporate organizations struggle with an inherent structural tension: centralization versus "
        "decentralization. Purely centralized operating models paralyze branch-level execution, forcing local plant managers to wait for "
        "headquarters approval to adjust daily shift schedules, handle customer SLA exceptions, or rebalance workstation labor. Conversely, "
        "purely decentralized models result in severe corporate anarchy: local branches manipulate pricing, delay financial journal reporting, "
        "obscure cash drawer balances, and violate corporate legal standards."
    )
    add_body_paragraph(doc, gov_p1, bold_prefix="The Centralization vs. Decentralization Dilemma: ")
    md_lines.append("**The Centralization vs. Decentralization Dilemma:** " + gov_p1 + "\n")

    gov_p2 = (
        "The German Hybrid Governance Architecture (modeled after advanced industrial engineering hierarchies in German manufacturing giants "
        "such as Siemens, Bosch, and BMW) resolves this dilemma by establishing a strict architectural separation of powers: Decentralized "
        "Operational Execution paired with Unconditional Centralized Financial and Legal Control. Under this model, branch laundries and "
        "plant production floors are granted full operational autonomy to manage intra-day workstation flow, touch sorting, automated WhatsApp "
        "customer verification, and local courier routing. Simultaneously, all financial ledgers, cash safe closes, B2B enterprise contracts, "
        "marketing cohort telemetry, and executive scorecards are strictly locked and governed by the C-Suite Super Admin Headquarters."
    )
    add_body_paragraph(doc, gov_p2, bold_prefix="The Hybrid Separation of Powers: ")
    md_lines.append("**The Hybrid Separation of Powers:** " + gov_p2 + "\n")

    add_heading_2(doc, "3.1 Detailed Breakdown of the 10 Synchronized Corporate Departments")
    md_lines.append("### 3.1 Detailed Breakdown of the 10 Synchronized Corporate Departments\n")

    dept_p1 = (
        "To operationalize this governance model, the entire MJRH software ecosystem—from the C-Suite Super Admin portal (`/admin`) down to "
        "individual branch workstation screens (`/stations/*`)—is structured into exactly 10 synchronized corporate departments. Configured "
        "verbatim in `components/app-sidebar.tsx` and governed by the RBAC permission engine in `hooks/use-auth.tsx`, these departments "
        "ensure that every corporate officer and blue-collar technician operates within strict, auditable boundaries without UI visual clutter "
        "or non-professional emoji symbols:"
    )
    add_body_paragraph(doc, dept_p1, bold_prefix="10 Synchronized Departments in Production Code: ")
    md_lines.append("**10 Synchronized Departments in Production Code:** " + dept_p1 + "\n")

    depts = [
        ("1. Vorstand (Executive Ownership & C-Suite Board)", 
         "Represents the supreme corporate headquarters of the SaaS company. Governs overall strategic direction, C-suite executive scorecards (`/executive`), live system video showcase and Fifth Settlement sales verification (`/demo`), and unconditional owner RBAC sovereignty. The Vorstand department retains exclusive override capabilities over corporate cap table structures, institutional pricing models, and enterprise valuation reporting."),
        ("2. Betriebsleitung (Production & Daily Plant Operations)", 
         "Controls the live industrial heart of the laundry enterprise. Governs the Chief Operating Officer dashboard (`/ops`), the daily operational command center (`/today`), and the 10 rotational workstation screens (`/stations/reception`, `/sorting`, `/cleaning`, `/drying-assembly`, `/ironing`, `/packing`, `/qc`, `/intake`, `/delivery`, `/cs`). Enforces touch hyper-automation and single-actor order ownership, ensuring that production flow never stalls due to bureaucratic bottlenecks."),
        ("3. Vertrieb (Sales, B2B Enterprise Deals, & Commercial Pipeline)", 
         "Manages commercial revenue generation and customer acquisition. Governs new order intake (`/orders/new`), unified order tracking (`/orders`), multi-parameter customer database (`/customers`), corporate CRM & loyalty programs (`/crm`), and B2B enterprise contract pipeline management (`/admin/biz-dev` linked directly to `enterprise_deals` table). Vertrieb officers utilize automated B2B quoting and contract generation tools to secure commercial hotel and hospital accounts."),
        ("4. Marketing & Analytik (Marketing Telemetry & BI)", 
         "Delivers advanced corporate business intelligence and campaign telemetry. Governs real-time marketing analytics (`/marketing`), peak intake/delivery hour visualization (10:00–12:00 AM peak tracking), campaign ROAS, Net Revenue Retention (NRR %), Churn rate %, SaaS Rule of 40 %, and EU GDPR / US CCPA PII anonymization toggles (`anonymizePII`). By analyzing customer ordering cadence and geographic density, Marketing & Analytik optimizes customer acquisition spend."),
        ("5. Kundenservice (Customer Care & SLA Governance)", 
         "Ensures institutional service quality and customer retention. Governs customer support dispatch (`/cs`), customer portal interface (`/customer-portal`), SLA breach monitoring, and formal exception dispute compensation and settlement workflows (`/customer-care` and `/stations/cs`). Kundenservice officers possess structured authority to issue instant digital refunds or reclean vouchers, resolving customer grievances within minutes."),
        ("6. Logistik & Fuhrpark (Transport, Logistics & Fleet Management)", 
         "Controls external dispatch, delivery routing, and courier operations. Governs live interactive GPS courier telemetry (`/live-map`), dedicated courier mobile workstations (`/driver`), delivery station handoffs (`/stations/delivery`), and automated algorithmic driver assignment without manual scanning. Logistik managers monitor courier turnaround times, route efficiency, and fuel consumption across multi-branch logistics networks."),
        ("7. Lager & Lieferkette (Warehouse, Inventory & Supply Chain)", 
         "Manages physical capital assets, consumables, and multi-branch infrastructure. Governs equipment maintenance logs, consumable chemical and spare part inventory (`/inventory`), branch enterprise configurations (`/branches`), and unit media storage buckets (`unit-media`). Lager officers establish automated reorder triggers for cleaning solvents and garment hangers, preventing production halts caused by stockouts."),
        ("8. Rechtsabteilung (Legal Affairs, Contract Vault & Dispute Litigation)", 
         "Safeguards corporate legal integrity and regulatory compliance. Governs the standalone legal department view (`/legal` and `/admin/legal`), persisting interactive records directly to `public.legal_contracts` and Supabase Storage (`legal-documents`). Manages Master Services Agreements (MSAs), vehicle operating licenses (`vehicle_license`), procurement contracts (`purchase_contract`), and active litigation claims (`litigation_dispute`)."),
        ("9. Personalwesen (Human Resources & Workforce Governance)", 
         "Manages employee lifecycle, payroll, and productivity verification. Governs Mawared HR time & attendance (`/staff/attendance`), Form 6 performance scorecards (`/staff/scorecard`), employee management with mandatory phone numbers and automated 1-click WhatsApp verification (`/staff/users`, `/staff/new`), daily shift schedules (`/staff/schedule`), leave requests, daily wages, and specialized ironing piece-rate payroll (`/staff/ironing-payroll`)."),
        ("10. Finanzwesen (Finance, Treasury & Double-Entry Accounting)", 
         "Serves as the central financial treasury and accounting command center. Governs general ledger accounting (`/finance`), automated double-entry bookkeeping (`/accounting`), journal audit logs (`/ledger`), accounts receivable collections (`/receivables`), one-movement cash safe closing (`/cash-closing`), corporate budget forecasting (`/budgets`), platform subscription billing (`/billing`), and institutional pricing settings (`/settings`).")
    ]

    for title, desc in depts:
        add_heading_3(doc, title)
        add_body_paragraph(doc, desc)
        md_lines.append(f"#### {title}\n{desc}\n")

    add_callout_box(doc, "GERMAN HYBRID GOVERNANCE DUE DILIGENCE SUMMARY", [
        "• Sovereign RBAC Enforcement: In hooks/use-auth.tsx, hasRole('owner') or isSuperAdmin grants unconditional modification authority across all 10 corporate departments.",
        "• Departmental Isolation: Floor technicians in Betriebsleitung cannot view executive treasury in Finanzwesen or litigation disputes in Rechtsabteilung, guaranteeing strict data security.",
        "• Zero-Emoji Professional UI: Every department heading, button label, and table header across all 10 screens relies strictly on clean Lucide icons and professional typography."
    ], border_color="1E293B", bg_color="F8FAFC")
    
    md_lines.append("> **GERMAN HYBRID GOVERNANCE DUE DILIGENCE SUMMARY**  \n"
                    "> * **Sovereign RBAC Enforcement:** In `hooks/use-auth.tsx`, `hasRole('owner')` or `isSuperAdmin` grants unconditional modification authority across all 10 departments.  \n"
                    "> * **Departmental Isolation:** Floor technicians in `Betriebsleitung` cannot view executive treasury in `Finanzwesen` or litigation disputes in `Rechtsabteilung`, guaranteeing strict data security.  \n"
                    "> * **Zero-Emoji Professional UI:** Every department heading, button label, and table header across all 10 screens relies strictly on clean Lucide icons and professional typography.\n")
    md_lines.append("\n---\n")

print("Sections 01 to 03 generated (Expanded).")
