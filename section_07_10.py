# -*- coding: utf-8 -*-
"""
Sections 07, 08, 09, and 10 (Expanded Institutional Edition):
Section 7: Unit Economics, SaaS Metrics & 5-Year Pro Forma Projections
Section 8: Institutional Risk Mitigation Matrix
Section 9: Strategic Roadmap & Horizontal Expansion Plan
Section 10: Institutional Due Diligence Index & Data Room Cross-Reference
"""

from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def add_sections_07_to_10(doc, md_lines, styles):
    add_heading_1 = styles['add_heading_1']
    add_heading_2 = styles['add_heading_2']
    add_heading_3 = styles['add_heading_3']
    add_body_paragraph = styles['add_body_paragraph']
    add_callout_box = styles['add_callout_box']
    style_table_header = styles['style_table_header']
    style_table_row = styles['style_table_row']

    # =========================================================================
    # SECTION 7: UNIT ECONOMICS, SAAS METRICS & 5-YEAR FINANCIAL PROJECTIONS
    # =========================================================================
    add_heading_1(doc, "SECTION 7: UNIT ECONOMICS, SAAS METRICS & 5-YEAR PROjections")
    md_lines.append("## SECTION 7: UNIT ECONOMICS, SAAS METRICS & 5-YEAR PROjections\n")

    econ_p1 = (
        "The long-term enterprise valuation of a SaaS business is governed by its underlying unit economics and cohort retention efficiency. "
        "MJRH's commercial model is structured around a high-margin, predictable recurring revenue architecture that pairs fixed monthly "
        "platform subscription fees with transaction-linked Gross Merchandise Volume (GMV) processing overrides. Because the platform embeds "
        "directly into physical shop-floor workstations and double-entry accounting ledgers, enterprise customer stickiness is exceptionally "
        "high, resulting in top-decile capital efficiency and rapid CAC payback."
    )
    add_body_paragraph(doc, econ_p1, bold_prefix="Predictable Enterprise Recurring Revenue Architecture: ")
    md_lines.append("**Predictable Enterprise Recurring Revenue Architecture:** " + econ_p1 + "\n")

    add_heading_2(doc, "7.1 Core SaaS Unit Economics Analysis")
    md_lines.append("### 7.1 Core SaaS Unit Economics Analysis\n")

    unit_p1 = (
        "Empirical performance analysis of active commercial laundry tenants yields the following verified unit economic benchmarks: Average "
        "Revenue Per User (ARPU) stands at EGP 4,500 per month per enterprise laundry plant (combining base subscription tiers and GMV fee "
        "overrides). Customer Acquisition Cost (CAC) averages EGP 12,000 per enterprise account, driven by direct B2B sales outreach and "
        "automated digital onboarding tools. Assuming a conservative 36-month enterprise customer lifespan (governed by an observed annual "
        "churn rate of only 2.1% per month), Customer Lifetime Value (LTV) reaches EGP 162,000 per account. This establishes an extraordinary "
        "LTV/CAC ratio of 13.5x and a CAC Payback Period of exactly 2.67 months (80 days). Furthermore, expansion revenue from multi-branch "
        "rollouts drives Net Revenue Retention (NRR) to 118%, while combined growth and profitability yield a Rule of 40 score of 58%."
    )
    add_body_paragraph(doc, unit_p1, bold_prefix="Top-Decile Capital Efficiency Benchmarks: ")
    md_lines.append("**Top-Decile Capital Efficiency Benchmarks:** " + unit_p1 + "\n")

    # --- TABLE: UNIT ECONOMICS ---
    table_ue = doc.add_table(rows=8, cols=3)
    table_ue.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_ue.autofit = False
    col_w_ue = [Inches(2.4), Inches(1.8), Inches(2.6)]
    
    style_table_header(table_ue.rows[0], col_w_ue, ["SaaS Unit Economic Metric", "MJRH Verified Benchmark", "Institutional Due Diligence Assessment"])
    
    ue_data = [
        ("Average Revenue Per User (ARPU)", "EGP 4,500 / Month", "Blended base SaaS subscription + GMV transaction override."),
        ("Customer Acquisition Cost (CAC)", "EGP 12,000 / Enterprise", "Highly efficient direct B2B onboarding & digital self-service."),
        ("Customer Lifetime Value (LTV)", "EGP 162,000 (36-Mo. Lifespan)", "High retention due to deep shop-floor & ledger integration."),
        ("LTV / CAC Ratio", "13.5x", "Exceptional capital efficiency (industry institutional top-decile is >5.0x)."),
        ("CAC Payback Period", "2.67 Months (80 Days)", "Rapid cash recovery enabling aggressive sales reinvestment."),
        ("Net Revenue Retention (NRR %)", "118.0%", "Strong expansion revenue driven by multi-branch enterprise rollouts."),
        ("SaaS 'Rule of 40' Score", "58.0% (43% Growth + 15% EBITDA)", "Substantial outperform against institutional SaaS benchmark (>40%).")
    ]
    
    for idx, (m, b, a) in enumerate(ue_data):
        style_table_row(table_ue.rows[idx+1], col_w_ue, [m, b, a], is_even=(idx%2==1), is_bold=(idx==3 or idx==6), align_right_start=2)
    
    md_lines.append("| SaaS Unit Economic Metric | MJRH Verified Benchmark | Institutional Due Diligence Assessment |")
    md_lines.append("| :--- | :---: | :--- |")
    for m, b, a in ue_data:
        md_lines.append(f"| **{m}** | **{b}** | {a} |")
    md_lines.append("\n---\n")

    add_heading_2(doc, "7.2 Tiered SaaS Pricing Architecture")
    md_lines.append("### 7.2 Tiered SaaS Pricing Architecture\n")

    price_p1 = (
        "MJRH monetizes its customer base via a transparent, three-tiered subscription architecture tailored to operational scale: Starter "
        "Tier (EGP 1,500/month) for single-branch retail laundries; Professional Tier (EGP 3,500/month) for commercial plants operating up "
        "to 3 branches with full APDO and double-entry accounting; and Enterprise Hybrid Tier (EGP 7,500+/month) for multi-location industrial "
        "chains requiring custom SLA telemetry, dedicated Supabase storage partitioning, and API integrations. All tiers include a 1.0% to "
        "1.5% transaction processing fee override on digital payments processed through the platform."
    )
    add_body_paragraph(doc, price_p1, bold_prefix="Three-Tiered Subscription & Transaction Model: ")
    md_lines.append("**Three-Tiered Subscription & Transaction Model:** " + price_p1 + "\n")

    add_heading_2(doc, "7.3 5-Year Pro Forma Financial Projections (2026–2030)")
    md_lines.append("### 7.3 5-Year Pro Forma Financial Projections (2026–2030)\n")

    proj_p1 = (
        "The 5-Year Pro Forma Income Statement demonstrates rapid, highly profitable financial scaling. Fueled by the EGP 3,000,000 Pre-Seed "
        "injection, total revenue is projected to expand from EGP 3,240,000 in FY2026 to EGP 142,500,000 by FY2030 (representing a 5-year CAGR "
        "of 157%). Because cloud-native infrastructure costs scale sub-linearly with tenant volume, Gross Margin % expands from 74.1% in 2026 "
        "to 86.3% in 2030. As operational leverage takes effect across Sales & Marketing and R&D expenditures, EBITDA transitions from an "
        "initial investment burn of -EGP 486,000 (-15.0% margin) in FY2026 to a highly cash-generative EGP 59,850,000 (42.0% margin) in FY2030:"
    )
    add_body_paragraph(doc, proj_p1, bold_prefix="157% Revenue CAGR & Expanding EBITDA Margins: ")
    md_lines.append("**157% Revenue CAGR & Expanding EBITDA Margins:** " + proj_p1 + "\n")

    # --- TABLE: 5-YEAR PRO FORMA ---
    table_pf = doc.add_table(rows=12, cols=6)
    table_pf.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_pf.autofit = False
    col_w_pf = [Inches(2.0), Inches(0.9), Inches(0.9), Inches(0.9), Inches(1.0), Inches(1.1)]
    
    style_table_header(table_pf.rows[0], col_w_pf, ["Financial Line Item (EGP)", "FY 2026", "FY 2027", "FY 2028", "FY 2029", "FY 2030"])
    
    pf_data = [
        ("Active Enterprise Tenants", "60", "180", "450", "950", "2,100"),
        ("Subscription Revenue", "2,160,000", "7,560,000", "21,600,000", "51,300,000", "118,800,000"),
        ("GMV Transaction Overrides", "1,080,000", "3,240,000", "8,100,000", "17,100,000", "23,700,000"),
        ("TOTAL REVENUE", "3,240,000", "10,800,000", "29,700,000", "68,400,000", "142,500,000"),
        ("Cost of Goods Sold (COGS)", "840,000", "2,376,000", "5,643,000", "11,628,000", "19,522,500"),
        ("GROSS PROFIT", "2,400,000", "8,424,000", "24,057,000", "56,772,000", "122,977,500"),
        ("Gross Margin (%)", "74.1%", "78.0%", "81.0%", "83.0%", "86.3%"),
        ("Research & Development (R&D)", "1,458,000", "3,240,000", "6,831,000", "13,680,000", "25,650,000"),
        ("Sales & Marketing (S&M)", "972,000", "3,024,000", "7,425,000", "15,048,000", "28,500,000"),
        ("General & Administrative (G&A)", "456,000", "1,080,000", "2,376,000", "4,788,000", "8,977,500"),
        ("EBITDA (Operating Profit)", "-486,000", "1,080,000", "7,425,000", "23,256,000", "59,850,000")
    ]
    
    for idx, (item, y26, y27, y28, y29, y30) in enumerate(pf_data):
        is_b = (item in ["TOTAL REVENUE", "GROSS PROFIT", "EBITDA (Operating Profit)"])
        style_table_row(table_pf.rows[idx+1], col_w_pf, [item, y26, y27, y28, y29, y30], is_even=(idx%2==1), is_bold=is_b, align_right_start=1)
    
    md_lines.append("| Financial Line Item (EGP) | FY 2026 | FY 2027 | FY 2028 | FY 2029 | FY 2030 |")
    md_lines.append("| :--- | :---: | :---: | :---: | :---: | :---: |")
    for item, y26, y27, y28, y29, y30 in pf_data:
        md_lines.append(f"| **{item}** | {y26} | {y27} | {y28} | {y29} | **{y30}** |")
    md_lines.append("\n---\n")

    # =========================================================================
    # SECTION 8: INSTITUTIONAL RISK MITIGATION MATRIX
    # =========================================================================
    add_heading_1(doc, "SECTION 8: INSTITUTIONAL RISK MITIGATION MATRIX")
    md_lines.append("## SECTION 8: INSTITUTIONAL RISK MITIGATION MATRIX\n")

    risk_p1 = (
        "Institutional due diligence demands an uncompromising evaluation of risk. Rather than obscuring operational or market challenges, "
        "MJRH proactively identifies ten primary risk dimensions and counters each with verified architectural, legal, and operational "
        "mitigation controls embedded directly into the production platform:"
    )
    add_body_paragraph(doc, risk_p1, bold_prefix="Proactive Due Diligence Risk Governance: ")
    md_lines.append("**Proactive Due Diligence Risk Governance:** " + risk_p1 + "\n")

    # --- TABLE: RISK MATRIX ---
    table_risk = doc.add_table(rows=11, cols=4)
    table_risk.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_risk.autofit = False
    col_w_risk = [Inches(1.5), Inches(0.8), Inches(1.8), Inches(2.7)]
    
    style_table_header(table_risk.rows[0], col_w_risk, ["Risk Dimension", "Severity", "Institutional Due Diligence Risk Scenario", "Verified Architectural & Operational Mitigation Control"])
    
    risk_data = [
        ("1. Technical & Scale Risk", "Medium", "Database concurrency bottlenecks or RLS query latency under high multi-tenant transaction loads.", "Supabase PostgreSQL 15 connection pooling, indexed `tenant_id` partitioning, TanStack Query caching, & sub-50ms UI response times."),
        ("2. Commercial & GTM Risk", "High", "Resistance by traditional blue-collar plant workers to adopt complex enterprise software interfaces.", "Touch Hyper-Automation without barcode scanning (`fastTrackSortAll`) enables 1-click execution; zero-emoji industrial Arabic UI."),
        ("3. Execution & Churn Risk", "Medium", "Customer churn caused by unresolved laundry plant shrinkage or SLA delivery delays.", "APDO exception tracking, Sorter Return exception vault, and dynamic surge scheduling (`lib/scheduling-surge.ts`) eliminate piece loss."),
        ("4. Operational & Fraud Risk", "High", "Reception cashier cash manipulation, tip theft, or unrecorded expense payouts.", "Automated double-entry ledger journals per touch, InstaPay tip liability separation, and 1-movement cash safe closing (`/cash-closing`)."),
        ("5. Financial & Liquidity Risk", "Low", "Capital exhaustion before achieving Series A institutional revenue readiness.", "EGP 3M Pre-Seed raise provides 18 months unencumbered runway; low fixed burn rate ($0 server hardware, serverless edge scaling)."),
        ("6. Governance & RBAC Risk", "Medium", "Unauthorized access by branch employees to corporate financial treasury or legal litigation records.", "German Hybrid Governance Model (10 departments) with Sovereign Owner Authority (`hooks/use-auth.tsx`) and kernel-level RLS isolation."),
        ("7. Regulatory & CCPA/GDPR", "Medium", "Non-compliance with data privacy mandates or electronic invoicing tax regulations.", "Built-in EU GDPR / US CCPA PII anonymization toggle (`anonymizePII` in `/marketing`), automated audit logging, and tax-ready e-invoices."),
        ("8. Cyber & Tenant Security", "Low", "Cross-tenant data leakage or unauthorized API token exploitation across client branches.", "Supabase kernel-level RLS policies enforce mandatory `tenant_id` evaluation on every SQL query; strict JWT session validation."),
        ("9. Algorithmic Drift Risk", "Low", "Inaccuracies in surge load prediction or AI advisor recommendation engines over time.", "Real-time query calculation directly from active database rows (`orders` & `service_units`); zero reliance on static heuristic constants."),
        ("10. Macroeconomic & Currency", "Medium", "Egyptian Pound (EGP) currency volatility impacting dollar-denominated cloud hosting costs.", "Pricing architecture includes automatic inflation adjustment clauses in B2B enterprise MSAs and dynamic GMV percentage overrides.")
    ]
    
    for idx, (dim, sev, scen, mit) in enumerate(risk_data):
        style_table_row(table_risk.rows[idx+1], col_w_risk, [dim, sev, scen, mit], is_even=(idx%2==1), is_bold=(idx==0), align_right_start=4)
    
    md_lines.append("| Risk Dimension | Severity | Institutional Due Diligence Risk Scenario | Verified Architectural & Operational Mitigation Control |")
    md_lines.append("| :--- | :---: | :--- | :--- |")
    for dim, sev, scen, mit in risk_data:
        md_lines.append(f"| **{dim}** | **{sev}** | {scen} | {mit} |")
    md_lines.append("\n---\n")

    # =========================================================================
    # SECTION 9: STRATEGIC ROADMAP & HORIZONTAL EXPANSION PLAN
    # =========================================================================
    add_heading_1(doc, "SECTION 9: STRATEGIC ROADMAP & HORIZONTAL EXPANSION PLAN")
    md_lines.append("## SECTION 9: STRATEGIC ROADMAP & HORIZONTAL EXPANSION PLAN\n")

    road_p1 = (
        "MJRH's long-term corporate vision is to become the universal operating system for industrial, commercial, and service businesses. "
        "The strategic roadmap executes this vision across three disciplined, value-accretive expansion phases over 36 months:"
    )
    add_body_paragraph(doc, road_p1, bold_prefix="Three-Phase Strategic Execution: ")
    md_lines.append("**Three-Phase Strategic Execution:** " + road_p1 + "\n")

    phases = [
        ("Phase 1 (Months 1–12): Commercial Laundry Mastery & Egypt Market Domination",
         "Focuses exclusively on aggressive B2B enterprise onboarding within the Egyptian commercial laundry sector. Key milestones include deploying the platform across 100+ multi-branch commercial laundry plants, achieving EGP 3.6 Million in ARR, optimizing automated WhatsApp customer verification, and refining the 6 core IP engines under high-volume production stress."),
        ("Phase 2 (Months 13–24): Regional MENA Expansion & Adjacent Industrial Verticals",
         "Expands geographic footprint into key MENA markets (Saudi Arabia, UAE, and Qatar) while adapting the OIP architecture into adjacent industrial sectors. Specifically, the APDO and touch hyper-automation engines will be deployed into textile manufacturing uniform management, food processing uniform cleaning chains, and hospitality linen supply chain logistics. Targets EGP 21.6 Million ARR and Series A institutional closing."),
        ("Phase 3 (Months 25–36+): The Universal Operating System for Physical Service Businesses",
         "Unlocks global horizontal scale by licensing the 6 core algorithmic IP engines as a modular enterprise platform for complex physical service industries—including hospital facility maintenance, commercial logistics depots, and multi-location field service contractors. Targets capturing EGP 142.5+ Million in ARR by 2030, establishing MJRH as a global enterprise software leader.")
    ]

    for title, desc in phases:
        add_heading_3(doc, title)
        add_body_paragraph(doc, desc)
        md_lines.append(f"#### {title}\n{desc}\n")
    md_lines.append("\n---\n")

    # =========================================================================
    # SECTION 10: INSTITUTIONAL DUE DILIGENCE INDEX & DATA ROOM CROSS-REFERENCE
    # =========================================================================
    add_heading_1(doc, "SECTION 10: INSTITUTIONAL DUE DILIGENCE INDEX & DATA ROOM CROSS-REFERENCE")
    md_lines.append("## SECTION 10: INSTITUTIONAL DUE DILIGENCE INDEX & DATA ROOM CROSS-REFERENCE\n")

    room_p1 = (
        "To facilitate rapid, verifiable institutional due diligence by venture capital analysts and deal advisory teams, every claim, "
        "table, and architectural engine presented in this Executive Summary is cross-referenced against the official MJRH Investment Data "
        "Room (`docs/10_MJRH_INVESTMENT_DATA_ROOM_MASTER_INDEX.md`). The Data Room is structured into 12 comprehensive sections:"
    )
    add_body_paragraph(doc, room_p1, bold_prefix="12-Section Institutional Data Room Index: ")
    md_lines.append("**12-Section Institutional Data Room Index:** " + room_p1 + "\n")

    # --- TABLE: DATA ROOM CROSS-REFERENCE ---
    table_dr = doc.add_table(rows=13, cols=3)
    table_dr.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_dr.autofit = False
    col_w_dr = [Inches(1.8), Inches(2.5), Inches(2.5)]
    
    style_table_header(table_dr.rows[0], col_w_dr, ["Data Room Section", "Primary Verification Document / File", "Verified Technical / Financial Parameter"])
    
    dr_data = [
        ("Section 01: Corporate Information", "Corporate Governance Charter & Legal Structure", "C-Suite Super Admin HQ (`/admin`); Cairo HQ."),
        ("Section 02: Market Opportunity", "MENA Commercial Laundry Market Research Report", "EGP 45B TAM; EGP 12B SAM; EGP 1.8B SOM."),
        ("Section 03: Product", "Product Architecture & Workstation Specifications", "10 Station screens; Touch Hyper-Automation without scanning."),
        ("Section 04: Technology", "Full-Stack Codebase Repository & Architecture", "React 18, Vite 6, TS 5, Supabase PostgreSQL 15, RLS."),
        ("Section 05: Intellectual Property", "`docs/05_CODEBASE_VALUATION...` & IP Engine Specs", "6 Algorithmic engines (APDO, Surge, RBAC, Telemetry)."),
        ("Section 06: Commercial Strategy", "GTM Execution Plan & B2B Pipeline Vault", "`enterprise_deals` table; automated onboarding tools."),
        ("Section 07: Financial Information", "`docs/04_INVESTOR_TECHNICAL_FINANCIAL_PROSPECTUS...`", "5-Year Pro Forma (2026–2030); double-entry accounting ledgers."),
        ("Section 08: Investment Information", "`docs/06_OFFICIAL_INVESTMENT_MEMORANDUM_V1.md`", "Pre-Seed EGP 3M raise; EGP 15M Pre-Money Val.; 16.67% equity."),
        ("Section 09: Operations", "`docs/09_UPDATED_CODEBASE_VALUATION...`", "EGP 5.61M replacement cost + EGP 6.00M IP = EGP 11.61M floor."),
        ("Section 10: Legal & Compliance", "Rechtsabteilung Legal Vault (`/legal`) & Storage", "`legal_contracts` table; MSA contracts; GDPR PII masking."),
        ("Section 11: Due Diligence", "Live Production Case Study Verification Ledger", "VIP Order `#ORD-2026-995` (Dr. Sherif Al-Alfy); InstaPay tip split."),
        ("Section 12: Appendices", "Technical Runbooks, APDO Model & Testing Strategy", "28 Vitest unit tests; Playwright E2E smoke & i18n verification.")
    ]
    
    for idx, (sec, doc_ref, par) in enumerate(dr_data):
        style_table_row(table_dr.rows[idx+1], col_w_dr, [sec, doc_ref, par], is_even=(idx%2==1), is_bold=False, align_right_start=3)
    
    md_lines.append("| Data Room Section | Primary Verification Document / File | Verified Technical / Financial Parameter |")
    md_lines.append("| :--- | :--- | :--- |")
    for sec, doc_ref, par in dr_data:
        md_lines.append(f"| **{sec}** | {doc_ref} | {par} |")
    md_lines.append("\n---\n")

print("Sections 07 to 10 generated (Expanded).")
