# -*- coding: utf-8 -*-
"""
Closing Statement & Contact Protocol
Glossary of Terms & Acronyms
"""

from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def add_closing_and_glossary(doc, md_lines, styles):
    add_heading_1 = styles['add_heading_1']
    add_heading_2 = styles['add_heading_2']
    add_heading_3 = styles['add_heading_3']
    add_body_paragraph = styles['add_body_paragraph']
    add_callout_box = styles['add_callout_box']
    style_table_header = styles['style_table_header']
    style_table_row = styles['style_table_row']

    # =========================================================================
    # CLOSING STATEMENT & CONTACT PROTOCOL
    # =========================================================================
    add_heading_1(doc, "CLOSING STATEMENT & INSTITUTIONAL CONTACT PROTOCOL")
    md_lines.append("## CLOSING STATEMENT & INSTITUTIONAL CONTACT PROTOCOL\n")

    close_p1 = (
        "MJRH INDUSTRIAL REVOLUTION represents a rare convergence of Category-Defining Enterprise Software Architecture, Deep Physical "
        "Industry Expertise, and Extraordinary Verified Asset Backing. By establishing the Operational Intelligence Platform (OIP) category, "
        "enforcing the German Hybrid Governance Architecture, and automating complex physical service workflows without scanning hardware, "
        "MJRH is uniquely positioned to dominate the EGP 45 Billion MENA commercial laundry sector and scale globally into industrial service "
        "operating systems. We invite institutional investors to review our comprehensive Data Room and schedule a formal investment committee "
        "presentation."
    )
    add_body_paragraph(doc, close_p1, bold_prefix="The Institutional Opportunity: ")
    md_lines.append("**The Institutional Opportunity:** " + close_p1 + "\n")

    add_callout_box(doc, "OFFICIAL CORPORATE CONTACT & DUE DILIGENCE ACCESS", [
        "• Authoring Executive: Mohamed Riyad — Founder & Chief Executive Officer",
        "• Corporate Headquarters: Cairo, Egypt | Operating Headquarters for MENA, EU, & US Expansion",
        "• Official Institutional Email: mohamedriyad1408@gmail.com",
        "• Live Platform Production Repository: https://mjrh.vercel.app/",
        "• Data Room Access Protocol: Institutional investors may request cryptographic access credentials to the 12-section Data Room by submitting a formal verification request to the official email above."
    ], border_color="0D9488", bg_color="F0FDFA")
    
    md_lines.append("> **OFFICIAL CORPORATE CONTACT & DUE DILIGENCE ACCESS**  \n"
                    "> * **Authoring Executive:** Mohamed Riyad — Founder & Chief Executive Officer  \n"
                    "> * **Corporate Headquarters:** Cairo, Egypt | Operating Headquarters for MENA, EU, & US Expansion  \n"
                    "> * **Official Institutional Email:** `mohamedriyad1408@gmail.com`  \n"
                    "> * **Live Platform Production Repository:** `https://mjrh.vercel.app/`  \n"
                    "> * **Data Room Access Protocol:** Institutional investors may request cryptographic access credentials to the 12-section Data Room by submitting a formal verification request to the official email above.\n")
    md_lines.append("\n---\n")

    # =========================================================================
    # GLOSSARY OF TERMS & ACRONYMS
    # =========================================================================
    add_heading_1(doc, "GLOSSARY OF TERMS & ACRONYMS")
    md_lines.append("## GLOSSARY OF TERMS & ACRONYMS\n")

    glossary_items = [
        ("OIP (Operational Intelligence Platform)", "An active cloud-native software architecture that digitizes, governs, measures, and optimizes physical operational execution across multi-station businesses."),
        ("APDO Operational Model", "The proprietary 4-stage state machine (Actor -> Process -> Data -> Output) guaranteeing data integrity and audit compliance across plant workflows."),
        ("German Hybrid Governance Model", "An organizational structure (Version 2.6 Hybrid) separating decentralized branch operational execution from strict centralized corporate financial treasury control across 10 departments."),
        ("RBAC (Role-Based Access Control)", "Security architecture enforcing granular permissions across 15 professional roles, granting absolute modification authority exclusively to Sovereign Owners (`owner` or `isSuperAdmin`)."),
        ("RLS (Row-Level Security)", "PostgreSQL database kernel security mechanism isolating tenant data (`tenant_id`) to prevent unauthorized cross-tenant data access."),
        ("Touch Hyper-Automation", "1-click workstation execution (`fastTrackSortAll`, `fastTrackPackAndReady`) enabling rapid batch processing and GPS dispatch without handheld barcode scanners."),
        ("Surge Scheduling Engine", "Dynamic capacity load balancing algorithm that caps booking slots and applies Express Surcharge pricing during peak intake hours (10:00 AM–12:00 PM)."),
        ("Double-Entry Accounting Ledger", "Automated bookkeeping system recording balanced debit and credit journal entries for every operational event, incorporating Tip Liability Separation and cash safe closing."),
        ("GDPR / CCPA PII Masking", "European and US data privacy compliance toggle (`anonymizePII` in `/marketing`) that encrypts personal customer PII on analytical dashboards."),
        ("TAM / SAM / SOM", "Total Addressable Market (EGP 45B MENA), Serviceable Addressable Market (EGP 12B), and Serviceable Obtainable Market (EGP 1.8B 3-year target)."),
        ("NRR (Net Revenue Retention)", "SaaS metric measuring recurring revenue growth from existing customers including upgrades and multi-branch expansions (MJRH benchmark: 118%)."),
        ("SaaS Rule of 40", "Institutional software benchmark where combined annual revenue growth rate % and EBITDA profit margin % exceed 40% (MJRH benchmark: 58%)."),
        ("OOXML (.docx)", "Office Open XML modern Microsoft Word document standard required for clean in-app rendering and institutional due diligence.")
    ]

    for term, defn in glossary_items:
        add_body_paragraph(doc, f"{defn}", bold_prefix=f"• {term}: ")
        md_lines.append(f"* **{term}:** {defn}")

print("Closing Statement and Glossary generated.")
